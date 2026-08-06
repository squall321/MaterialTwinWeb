#!/usr/bin/env python
# 적색발광 액자형수축 과제 — 시뮬레이션 물성·확보현황·출처를 docx 한 부로 묶는다.
# 사용: .venv/bin/python scripts/catalog/build_redshrink_docx.py
from __future__ import annotations

import json
import sqlite3
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

DB = "/home/koopark/claude/HEAXHub/var/app_data/materialtwin_web/materialtwin.db"
OUT = Path("/home/koopark/claude/MaterialTwinWeb/적색발광_액자형수축_시뮬레이션물성_출처정리.docx")
BARRER = 3.3464e-16  # mol*m/(m^2*s*Pa)

# 해석 단계 → (물성키, 표기, 단위, 왜 필요한가)
STAGES = [
    ("산소 확산장 — 에지에서 중앙으로", [
        ("physical.diffusion_coefficient", "O₂ 확산계수 D", "m²/s",
         "액자 폭의 시간 법칙 w(t) ∝ √(D·t)를 정한다"),
        ("physical.gas_solubility", "헨리 용해도 S", "mol/(m³·Pa)",
         "외기 산소 분압 → 막 내부 용존 농도의 경계조건 c = S·pO₂"),
        ("physical.gas_permeability_o2", "산소 투과도 P", "mol/(m·s·Pa)",
         "P = D·S. 에지 경로와 면방향의 상대 크기를 정한다"),
        ("chemical.activation_energy", "활성화에너지 Ea", "J/mol",
         "온도 가속 계수 — 시험계획 Phase 2가 온도 3수준을 도는 근거"),
    ]),
    ("Stern-Volmer 소광", [
        ("optical.excited_state_lifetime", "여기수명 τ₀", "s",
         "I₀/I = 1 + k_q·τ₀·[O₂]. R/G/B 선택성을 결정하는 핵심 항"),
        ("optical.stern_volmer_constant", "Stern-Volmer 상수 K_SV", "1/Pa",
         "K_SV = k_q·τ₀. 분압에서 소광비를 직접 준다"),
        ("optical.bimolecular_quenching_rate", "소광 속도상수 k_q", "m³/(mol·s)",
         "소광의 절대 크기. τ₀와 곱해 K_SV가 된다"),
        ("optical.photoluminescence_quantum_yield", "광발광 양자수율 PLQY", "1",
         "탈기/공기 쌍이 있으면 산소 소광이 직접 관측된다"),
    ]),
    ("UV 리셋 — 잔류 개시제 라디칼의 산소 소거", [
        ("optical.molar_absorptivity", "몰흡광계수 ε", "m²/mol",
         "개시제 흡수대와 UV 액션 스펙트럼의 겹침 — H1c의 판별 지문"),
    ]),
    ("UV·외광의 내부 도달 경로", [
        ("optical.transmittance", "분광 투과율 T", "1",
         "적색 CF의 380~420 nm 2차 투과창 — 외광이 적색 화소에 닿는 경로"),
    ]),
    ("비가역 전이 — 가역 구간의 시한", [
        ("optical.singlet_oxygen_quantum_yield", "일중항 산소 수율 Φ_Δ", "1",
         "소광 에너지가 ¹O₂로 넘어가 영구 광산화로 전이하는 비율"),
    ]),
]

# 4단계는 CF·개시제만 담는다(카탈로그 전체 투과율은 이 과제와 무관).
CF_ONLY = {"optical.transmittance"}


def _sty(doc):
    n = doc.styles["Normal"]
    n.font.name = "맑은 고딕"
    n.font.size = Pt(9.5)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    for i, sz in ((1, 15), (2, 12.5), (3, 11)):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "맑은 고딕"
        h.font.size = Pt(sz)
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        h.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _tbl(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(8)
    return t


def _p(doc, text, bold=False, size=9.5, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


def fmt(v, key):
    if v is None:
        return "—"
    if key == "optical.excited_state_lifetime":
        return f"{v*1e9:.3g} ns" if v < 1e-7 else f"{v*1e6:.4g} µs"
    if key == "physical.gas_permeability_o2":
        return f"{v/BARRER:.4g} Barrer"
    if key == "chemical.activation_energy":
        return f"{v/1000:.4g} kJ/mol"
    if abs(v) >= 1e4 or (v != 0 and abs(v) < 1e-3):
        return f"{v:.4g}"
    return f"{v:.6g}".rstrip("0").rstrip(".") if "." in f"{v:.6g}" else f"{v:.6g}"


def cond_brief(cond, keys, limit=110):
    if not cond:
        return ""
    try:
        d = json.loads(cond)
    except Exception:
        return ""
    bits = []
    for k in keys:
        if d.get(k) not in (None, ""):
            bits.append(f"{k}={d[k]}")
    rest = [f"{k}={v}" for k, v in d.items()
            if k not in keys and v not in (None, "") and not isinstance(v, (dict, list))]
    s = ", ".join(bits + rest)
    return s[:limit] + ("…" if len(s) > limit else "")


def src_label(row):
    """출처를 한 줄로. 저널이면 DOI, 데이터시트면 발행처."""
    title, kind, doi, url, pub = row
    t = (title or "").strip()
    tag = doi or (pub or "") or (url or "")
    if doi:
        tag = f"DOI {doi}"
    elif url:
        tag = url
    return t[:150], (tag or "")[:110], kind or ""


def main():
    c = sqlite3.connect(DB)
    doc = Document()
    _sty(doc)

    # ── 표지 ──────────────────────────────────────────────────────────────
    h = doc.add_heading("적색 발광 액자형 수축", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("반응·확산 시뮬레이션 물성 · 확보 현황 · 출처")
    r.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n_mat, n_val, n_src, n_def = c.execute(
        "select (select count(*) from material),(select count(*) from property_value),"
        "(select count(*) from source),(select count(*) from property_definition)").fetchone()
    r = p.add_run(f"2026-08-06 · MaterialTwin 카탈로그 "
                  f"재료 {n_mat:,}종 · 물성값 {n_val:,}건 · 출처 {n_src:,}건 · 물성 정의 {n_def}종")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    _p(doc, "이 문서는 시험계획서(산소영향성 시험계획)가 요구하는 산소 확산 × Stern-Volmer 소광 "
            "모델의 입력 물성을 해석 단계별로 정리하고, 확보한 값과 그 출처를 전부 수록한다. "
            "수록된 값은 모두 원문에 인쇄된 수치이며, 그래프에서 읽은 값·역산한 값·문헌 취합표의 "
            "재인용은 포함하지 않는다.")

    # ── 1. 결론 ───────────────────────────────────────────────────────────
    doc.add_heading("1. 결론 — 지금 되는 것과 안 되는 것", level=1)
    _p(doc, "시험 조건 설계용 프로토타입은 지금 돌릴 수 있다. 설계 판단(목표 수명 역산)은 못 한다.",
       bold=True, size=10.5)
    _p(doc, "이 구분은 시험계획이 스스로 정한 순서와 일치한다 — "
            "“프로토타입을 시험 착수 전에 완성해 시험 조건 설계에 활용한다.”")

    _p(doc, "지금 되는 것", bold=True, space=2)
    _tbl(doc, ["항목", "근거"], [
        ["가설 검증식 — R/G/B 폭 비율이 τ₀ 비율을 따르는가", "τ₀ 34종 / 73건. 적·녹·청 전부 확보"],
        ["UV 리셋 기전 판별 — 개시제 계열로 성립/불성립이 갈린다", "몰흡광계수 6종 / 16건"],
        ["UV·외광 도달 경로 — 적색 CF의 400 nm 창", "분광 투과율 실측 확보"],
        ["기하 논거 — 왜 4변 대칭 액자인가", "에지 경로 vs 면방향에 약 9자리 차"],
        ["파라미터 스윕 — 시간 법칙 형태, 온도 가속 계수 범위", "투과도 48종 / 104건, Ea 14종 / 26건"],
    ])
    doc.add_paragraph()

    _p(doc, "지금 안 되는 것", bold=True, space=2)
    _tbl(doc, ["못 하는 것", "없는 물성"], [
        ["절대 시간 스케일 — “3개월이면 액자 폭 몇 mm”", "실제 OCA의 O₂ 확산계수 D · 용해도 S가 0건"],
        ["R/G/B 절대 소광량 (비율은 가능)", "OLED 도판트 계열의 k_q가 0건"],
        ["가역 구간의 시한 — 언제부터 회수 불능인가", "목표 계열의 일중항 산소 수율 Φ_Δ가 0건"],
    ])
    doc.add_paragraph()
    _p(doc, "단서 — D가 없다는 것은 시뮬레이션이 못 돈다는 뜻이 아니다. D는 원래 캘리브레이션 "
            "대상이다. 시험계획 Phase 2(온도 3수준 × 시간 5시점)가 w(t) ∝ √(D_eff·t)에서 "
            "D_eff를 역피팅하는 설계이고, 문서도 “미지수가 사실상 두 개로 줄고 나머지는 문헌치와 "
            "별도 측정으로 고정된다”고 썼다. 시험 전에 절대값을 못 낼 뿐, 모델 골격과 시험 조건 "
            "설계는 지금 가능하다.", bold=True)

    # ── 2. 시험계획에 반영할 것 ────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("2. 데이터가 시험계획에 대해 말한 것", level=1)

    doc.add_heading("2.1  τ₀ 순서는 R > G > B 가 아니라 R > B > G 다", level=2)
    _p(doc, "같은 논문·같은 프로토콜(탈기 CH₂Cl₂) 안에서만 비교한 값이다.")
    _tbl(doc, ["", "τ₀", "녹색 대비"], [
        ["녹색 인광 Ir(ppy)₂(acac)", "0.73 µs", "1.00×"],
        ["청색 인광 FIrpic", "1.82 µs", "2.49×"],
        ["적색 인광 piq 계열", "3.90 ~ 11.2 µs", "5.3 ~ 15.3×"],
        ["청색 형광 (단일지수) pSFIAc1/2", "7.7 ~ 8.0 ns", "—"],
    ])
    doc.add_paragraph()
    _p(doc, "적/청 판별력", bold=True, space=2)
    _tbl(doc, ["청색 기준", "배율"], [
        ["청색 인광 FIrpic 1.82 µs", "2.1 ~ 6.2배"],
        ["청색 형광 1.6 ~ 13 ns", "300 ~ 7,000배"],
        ["순수 형광(단일지수) 7.7 ~ 8.0 ns", "488 ~ 1,455배"],
    ])
    doc.add_paragraph()
    _p(doc, "보고서 문구 확인 — 임원보고의 “적·녹은 발광 지속시간이 청색보다 1000배 가까이 길어”는 "
            "청색이 형광일 때 488~1,455배로 정확하다. 다만 이것은 유도된 결과가 아니라 전제다. "
            "청색 인광 스택이면 2~6배로 떨어져 판별력이 거의 사라지므로, 문서에 전제를 명시할 것.",
       bold=True)

    doc.add_heading("2.2  τ₀만으로는 소광량을 예측할 수 없다 — 매질이 그만큼 지배한다", level=2)
    _p(doc, "같은 PtOEP를 매질만 바꿔 잰 값이다.")
    _tbl(doc, ["매질", "τ₀ (진공)", "공기 중", "소광비"], [
        ["PCL 나노섬유", "60.5 µs", "6.2 µs", "9.76×"],
        ["PU Tecophilic", "80.8 µs", "11.1 µs", "7.28×"],
        ["PS 나노섬유", "83.9 µs", "15.2 µs", "5.52×"],
        ["PVDF-HFP", "81.4 µs", "62.0 µs", "1.31×"],
    ])
    doc.add_paragraph()
    _p(doc, "τ₀가 80 µs로 거의 같은데 소광비가 7배 갈린다. 시험 판정식에 반영할 것 — Phase 1의 "
            "R/G/B 폭 비율이 τ₀ 비율과 어긋나도 곧바로 기각이 아니다. 적·녹·청 발광층의 호스트가 "
            "다르면 매질 항이 개입한다.", bold=True)

    doc.add_heading("2.3  개시제 계열 하나로 가설이 갈린다", level=2)
    _tbl(doc, ["계열", "385~405 nm 몰흡광계수"], [
        ["아실포스핀옥사이드 TPO · BAPO(819)", "51 ~ 66 m²/mol"],
        ["하이드록시알킬페논 184 · 1173 · 2959 계열", "0.2 ~ 0.4 m²/mol"],
    ])
    doc.add_paragraph()
    _p(doc, "170 ~ 330배 차이다. OCA에 아실포스핀옥사이드계가 들어 있으면 H1c의 UV 리셋 기전이 "
            "성립하고, 하이드록시알킬페논계만 있으면 성립하지 않는다. 시험계획 §11의 "
            "“OCA 광개시제 계열” 업체 요청이 왜 결정적인지가 이것으로 확정된다 — "
            "공문 회신 하나로 Phase 1 착수 전에 가설을 걸러낼 수 있다.", bold=True)
    _p(doc, "부수 예측 — TPO 흡수가 395 → 420 nm에서 1/25로 급락한다. 따라서 UV 액션 스펙트럼의 "
            "장파장 끝이 415~420 nm에 걸려야 한다. 시험이 이를 확인하면 H1c의 강한 증거다.")

    doc.add_heading("2.4  적색 CF의 2차 투과창은 실재하고 의도된 것이다", level=2)
    _p(doc, "실제 경화막(0.9~1.0 µm) 실측이다.")
    _tbl(doc, ["시료", "T400", "T530 (자기 저지대)", "T610 (통과대역)"], [
        ["Example 14 (1.0 µm)", "10 %", "9 %", "99 %"],
        ["Example 23 (1.0 µm)", "17 %", "11 %", "70 %"],
        ["Example 24 (0.9 µm)", "16 %", "11 %", "72 %"],
    ])
    doc.add_paragraph()
    _p(doc, "세 시료 모두 T400 > T530 이다 — 잔류 누설이 아니라 별도의 창이다. T410은 다른 특허의 "
            "청구항이 0.4 < T410 < 4 를 요구한다. 색도·콘트라스트를 맞추려고 의도적으로 열어 둔 "
            "창이지 사고가 아니다. 이 창이 개시제 흡수대(TPO 380 nm 52~57 m²/mol)와 정확히 겹친다.")

    doc.add_heading("2.5  가교밀도 — 부호가 뒤집히는 계가 있다", level=2)
    _p(doc, "가교제 함량 조건이 붙은 산소투과도 53건 / 17재료. 11계 중 9계는 가설 방향이지만 2계가 반대다.")
    _tbl(doc, ["계", "가교제", "P (Barrer)", "방향"], [
        ["PEGDA / 2-EEA", "94.7 → 35 wt%", "5.3 → 19", "가설대로 ↑"],
        ["BPAEDA4 / PEGMEA", "85 → 28 wt%", "1.6 → 14", "가설대로 ↑ (8.8배)"],
        ["PEGDA / 2-HEA", "96 → 40 wt%", "4.4 → 0.45", "반대 ↓ (10배 감소)"],
        ["PEGDA / DEGPEA", "92 → 0 wt%", "4.4 → 2.9", "반대 ↓"],
    ])
    doc.add_paragraph()
    _p(doc, "하이드록실·페닐 공단량체가 사슬 패킹을 조여 가교밀도 효과를 뒤집는다. "
            "가교밀도만으로 OCA를 외삽하면 안 된다 — 공단량체 화학이 부호를 바꾼다. "
            "비가교(0 wt%) 하한점 2건(2.9 · 7.0 Barrer)을 확보했다. "
            "“미경화 OCA가 고속 통로”를 직접 시험할 수 있는 유일한 0점이다.", bold=True)

    # ── 3. 단계별 물성 요구와 확보 ──────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("3. 해석 단계별 물성 요구와 확보 현황", level=1)

    counts = {}
    for _, items in STAGES:
        for key, *_ in items:
            if key == "physical.diffusion_coefficient":
                m, v = c.execute(
                    "select count(distinct material_id),count(*) from property_value "
                    "where property_key=? and conditions like '%\"O2\"%'", (key,)).fetchone()
            elif key in CF_ONLY:
                m, v = c.execute(
                    "select count(distinct p.material_id),count(*) from property_value p "
                    "join material m on m.id=p.material_id where p.property_key=? "
                    "and (m.name like '%olor filter%' or m.name like '%olour filter%')", (key,)).fetchone()
            else:
                m, v = c.execute(
                    "select count(distinct material_id),count(*) from property_value "
                    "where property_key=?", (key,)).fetchone()
            counts[key] = (m, v)

    for si, (stage, items) in enumerate(STAGES, start=1):
        doc.add_heading(f"3.{si}  {stage}", level=2)
        rows = []
        for key, label, unit, why in items:
            m, v = counts[key]
            state = f"{m}종 / {v}건" if v else "0건"
            rows.append([label, unit, state, why])
        _tbl(doc, ["물성", "단위", "확보", "왜 필요한가"], rows)
        doc.add_paragraph()

    # 에지 vs 면방향 사다리
    doc.add_heading("3.6  에지 경로 vs 면방향 — 9자리 차이", level=2)
    _p(doc, "1 Barrer = 3.3464e-16 mol·m/(m²·s·Pa)")
    ladder = [("TFE ALD Al₂O₃ 38 nm", 6.5e-8), ("Toppan GX-P-F 배리어필름", 1.8e-5),
              ("Upilex-S", 3.1e-4), ("PET", 0.029), ("Kapton HN", 0.147),
              ("PC", 0.745), ("아크릴레이트(고무상) 하한", 19.0)]
    _tbl(doc, ["재료", "산소 투과도 (Barrer)"], [[n, f"{v:.4g}"] for n, v in ladder])
    doc.add_paragraph()
    _p(doc, "TFE 면방향과 OCA 에지 경로 사이가 약 9자리 벌어진다. "
            "“산소는 에지로만 들어온다”는 4변 대칭 논거의 정량 근거가 이것이다.", bold=True)

    # ── 4. 확보 물성 전량 + 출처 ────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("4. 확보한 물성값 전량과 출처", level=1)
    _p(doc, "값은 전부 원문에 인쇄된 수치다. tier 1 = 그 재료를 직접 측정해 인쇄한 값, "
            "tier 2 = 핸드북·공인 DB, tier 3 = 클래스 대표값·2차 인용, tier 4 = 가정·추정값.")

    COND_KEYS = ["matrix", "atmosphere", "term", "quencher", "species", "wavelength_nm",
                 "temperature_C", "temperature_K", "temperature_k", "cure_state",
                 "crosslinker_wt_pct", "solvent", "color", "thickness_um", "state"]

    for si, (stage, items) in enumerate(STAGES, start=1):
        doc.add_heading(f"4.{si}  {stage}", level=2)
        for key, label, unit, _why in items:
            if key == "physical.diffusion_coefficient":
                q = ("select m.name,p.value_num,p.quality_tier,p.conditions,p.source_id "
                     "from property_value p join material m on m.id=p.material_id "
                     "where p.property_key=? and p.conditions like '%\"O2\"%' order by m.name")
            elif key in CF_ONLY:
                q = ("select m.name,p.value_num,p.quality_tier,p.conditions,p.source_id "
                     "from property_value p join material m on m.id=p.material_id "
                     "where p.property_key=? and (m.name like '%olor filter%' "
                     "or m.name like '%olour filter%') order by m.name, p.id")
            else:
                q = ("select m.name,p.value_num,p.quality_tier,p.conditions,p.source_id "
                     "from property_value p join material m on m.id=p.material_id "
                     "where p.property_key=? order by m.name, p.id")
            rows = c.execute(q, (key,)).fetchall()
            if not rows:
                doc.add_heading(f"{label}  —  0건", level=3)
                _p(doc, "확보하지 못했다. 사유는 5장 참조.")
                continue

            doc.add_heading(f"{label}  ({unit})  —  {len(rows)}건", level=3)

            # 출처를 번호로 묶는다
            srcs, order = {}, []
            for _, _, _, _, sid in rows:
                if sid and sid not in srcs:
                    srcs[sid] = len(order) + 1
                    order.append(sid)

            body = []
            for nm, v, t, cond, sid in rows:
                body.append([nm[:52], fmt(v, key), f"t{t}",
                             cond_brief(cond, COND_KEYS), f"[{srcs.get(sid, '-')}]"])
            _tbl(doc, ["재료", "값", "등급", "조건", "출처"], body)
            doc.add_paragraph()

            _p(doc, "출처", bold=True, space=2)
            sr = []
            for sid in order:
                row = c.execute("select title,kind,doi,url,publisher from source where id=?",
                                (sid,)).fetchone()
                if not row:
                    continue
                title, tag, kind = src_label(row)
                sr.append([f"[{srcs[sid]}]", title, tag, kind])
            _tbl(doc, ["#", "제목", "DOI / URL", "종류"], sr)
            doc.add_paragraph()

    # ── 5. 미확보와 조달 경로 ──────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("5. 미확보 물성과 조달 경로", level=1)

    doc.add_heading("5.1  1순위 — 절대 시간 스케일을 여는 것", level=2)
    _tbl(doc, ["필요", "왜", "상태"], [
        ["실제 OCA의 O₂ 확산계수 D", "없으면 “몇 개월에 몇 mm”를 못 답한다", "문헌으로 채울 수 없음"],
        ["아크릴레이트 O₂ 용해도 S", "외기 분압 → 막 내부 농도 경계조건", "문헌으로 채울 수 없음"],
    ])
    doc.add_paragraph()
    _p(doc, "이 둘은 문헌으로 못 채운다. \"optically clear adhesive\" AND \"Barrer\" 검색 0건이고 "
            "3M OCA 데이터시트에 OTR 항목 자체가 없다. OCA는 배리어 부품이 아니라 업체가 측정·공개할 "
            "동기가 없다. 용해도는 원문이 사유를 인쇄한다 — 비극성 영구기체는 극성 고무에서 용해도가 "
            "낮아 압력감쇠법 검출한계에 걸린다. 실제로 어느 표에서도 CH₄·CO₂·C₂H₄·C₂H₆는 확산·용해도가 "
            "다 채워져 있는데 O₂·N₂·H₂·He만 빈칸이다.")
    _p(doc, "실질적 조달 경로는 셋뿐이다.", bold=True, space=2)
    _tbl(doc, ["#", "경로", "비고"], [
        ["1", "Phase 2 데이터로 D_eff 역피팅", "시험계획이 이미 그렇게 설계돼 있다 — 1순위"],
        ["2", "OCA 필름 단품 시편으로 OTR 외주 측정", "§11 업체 요청에 시편 항목이 이미 있다"],
        ["3", "대리계 범위(19 ~ 140 Barrer)로 파라미터 스윕", "절대값 대신 상하한을 낸다"],
    ])
    doc.add_paragraph()

    doc.add_heading("5.2  2순위 — 시험 판정의 정밀도", level=2)
    _tbl(doc, ["필요", "현황"], [
        ["OLED 도판트 계열의 k_q", "지금 2건인데 둘 다 Ir-쿠마린 센서 화합물이고 매질이 CH₂Cl₂ 용액이다. 발광층 고체 매질 값이 필요하다"],
        ["박막·고체 매질의 τ₀", "확보한 73건 대부분이 용액이다. 2.2절이 보여주듯 매질이 소광을 7배 좌우한다"],
        ["목표 계열(Ir·Pt OLED 도판트)의 Φ_Δ", "가역 구간의 시한을 정하는 항. 확보한 13건은 포르피린·Pt 착체다"],
        ["상용 청색 도판트 τ₀ (DPAVBi·BCzVBi·ADN·TBPe)", "OA에 소자 특성표만 있고 용액 광물리 측정이 없다. 지금은 같은 파장대 최신 발광체로 대리한다"],
    ])
    doc.add_paragraph()

    doc.add_heading("5.3  3순위 — 광경로 정밀화", level=2)
    _tbl(doc, ["필요", "현황"], [
        ["적색 CF 380 · 390 nm 투과율", "0건 — 창의 단파장 쪽 형상을 모른다"],
        ["청색 CF 400 nm 실측", "0건 — 청구항 범위(50~70%)만 있다. 청색이 UV 유입의 가장 큰 문일 가능성이 높다"],
        ["녹색 CF 도막의 400 nm", "0건 — 안료 분산액 값만 있다"],
        ["Irgacure 184 · 651 · 1173의 ε", "0건 — 벤더 PDS가 전부 스펙트럼 그림이고 논문은 정성 서술뿐이다"],
    ])
    doc.add_paragraph()

    doc.add_heading("5.4  채울 수 없는 것 — 구조적 부재", level=2)
    _tbl(doc, ["항목", "이유"], [
        ["TFE의 정밀 산소 투과도",
         "요구사양(1e-5 cm³/m²/day)이 상용 계측기 하한(약 5e-3) 아래라 업계가 산소는 아예 안 재고 "
         "수분(WVTR)만 보고한다. 문헌값이 전부 부등호 표기다"],
        ["아크릴 OCA의 산소 실측", "5.1절 참조"],
        ["poly(n-butyl acrylate) C3·C4 투과도",
         "원표(Mogri & Paul, Polymer 2001)가 유료. 저자 학위논문은 전자 납본 이전이라 리포지토리에 없다"],
    ])

    # ── 6. 수집 규율 ──────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("6. 수집 규율 — 무엇을 걸러냈나", level=1)
    _p(doc, "이 과제의 문헌은 값을 그림에만 두는 관행이 특히 강해, 그래프 판독 금지가 가장 크게 작동했다.")
    _tbl(doc, ["폐기 사유", "내용"], [
        ["그래프만 있음 (10건)", "분광·개시제 스펙트럼. 값이 Figure에만 있고 표에 없다"],
        ["특허 청구항 범위값", "설계 상한이지 측정값이 아니다. 다만 실시예 표의 실측은 등록했다 — 이 과제 성과의 상당수가 거기서 나왔다"],
        ["문헌 취합표", "Nature Communications의 표준 도판트 14종 τ 표가 매질·분위기 미표기라 통째로 폐기. 2.2절이 보여주듯 매질 없는 τ₀는 60배까지 흔들린다"],
        ["MD 시뮬레이션", "Kapton 모사 O₂ 투과도가 실측 대비 3000배 어긋났다"],
        ["물리적으로 불가능한 표", "PDMS가 N₂ > O₂ > CO₂로 나오는 표. 추출 오류 또는 원문 오식"],
        ["검출 하한 부등호 표기 (7건)", "“< 1e-3 g/m²/day” 같은 상한 표기는 실측이 아니다"],
        ["옆 칸 오독 회피", "K_SV(M⁻¹) 열의 소광종이 O₂가 아니라 TTA-UC 소멸제 DPA였다"],
        ["클래스 항목에 값 미부착", "OLED Emitter Layer (Ir phosphor)는 0건을 유지한다. PtOEP 60 µs와 Ir 계열 1~11 µs가 60배 차이라, 클래스에 특정 화합물 값을 붙이면 대표값을 지어내는 것이 된다"],
    ])
    doc.add_paragraph()

    doc.add_heading("6.1  환산 상수 오류를 수집 배치가 잡았다", level=2)
    _p(doc, "작업 지시서에 1 Barrer = 3.348e-19 이라고 적었는데 1000배 틀린 값이다. "
            "정확한 값은 3.3464e-16 mol·m/(m²·s·Pa) 이고 네 방법으로 확증됐다.")
    _tbl(doc, ["#", "확증 방법", "결과"], [
        ["1", "정의에서 직접 유도", "1e-10 × (1/22414) × 1e-2 / (1e-4 × 1333.2239) = 3.3464e-16"],
        ["2", "카탈로그 기존 실측 역검산", "PET 0.0289 Barrer(문헌 ~0.03), Kapton 0.147(~0.15), PC 0.745(~0.8)"],
        ["3", "SI·Barrer 병기 문헌", "Membranes 2022, 12, 611 — 4.79e-17 / 1.434e-1 = 3.3427e-16 (0.2% 일치)"],
        ["4", "리뷰가 상수를 직접 인쇄", "Macromolecules 2021 — 3.348e-16"],
    ])
    doc.add_paragraph()
    _p(doc, "수집 배치가 지시서를 그대로 믿지 않고 스스로 검산한 덕에 1000배 틀린 값이 대량 유입되지 않았다.",
       bold=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"저장: {OUT}")
    print(f"  단계 {len(STAGES)}개 · 수록 물성 {len(counts)}종")


if __name__ == "__main__":
    main()
