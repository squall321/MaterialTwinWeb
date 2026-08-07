#!/usr/bin/env python
# 카탈로그 전체 물성 확보 현황 · 해석 가능성 · 공백을 docx 한 부로 묶는다.
# 사용: .venv/bin/python scripts/catalog/build_coverage_docx.py
from __future__ import annotations

import sqlite3
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from coverage_report import ANALYSES, compute, property_stats  # noqa: E402

OUT = Path("/home/koopark/claude/MaterialTwinWeb/MaterialTwin_물성확보현황_해석가능성.docx")
DB = "/home/koopark/claude/HEAXHub/var/app_data/materialtwin_web/materialtwin.db"

DOMAIN_KO = {"mechanical": "기계", "thermal": "열", "electrical": "전기", "optical": "광학",
             "physical": "물리", "chemical": "화학", "interface": "계면", "structure": "구조",
             "rheological": "유변", "magnetic": "자기", "acoustic": "음향"}


def _sty(doc):
    n = doc.styles["Normal"]
    n.font.name = "맑은 고딕"
    n.font.size = Pt(9.5)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    for i, sz in ((1, 15), (2, 12.5), (3, 11)):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "맑은 고딕"
        h.font.size = Pt(sz)
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        h.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _tbl(doc, headers, rows, hsz=8.5, bsz=8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(hsz)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(bsz)
    return t


def _p(doc, text, bold=False, size=9.5, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


def bar(pct, width=20):
    """텍스트 막대 — 표에서 한눈에 비교되게."""
    n = int(round(pct / 100 * width))
    return "█" * n + "·" * (width - n)


def main():
    c, mat, own, cov = compute()
    props = property_stats(c, mat, own)
    doc = Document()
    _sty(doc)

    n_mat, n_val, n_src, n_def = c.execute(
        "select (select count(*) from material),(select count(*) from property_value),"
        "(select count(*) from source),(select count(*) from property_definition)").fetchone()
    tot_cells = sum(x["cells"] for x in cov)
    tot_filled = sum(x["filled"] for x in cov)
    tot_pct = tot_filled * 100 / tot_cells

    # ── 표지 ─────────────────────────────────────────────────────────────
    h = doc.add_heading("MaterialTwin 물성 확보 현황", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("해석별 커버리지 · 무엇이 채워졌고 무엇이 비었나")
    r.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"2026-08-06 · 재료 {n_mat:,}종 · 물성값 {n_val:,}건 · "
                  f"출처 {n_src:,}건 · 물성 정의 {n_def}종")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()

    # ── 1. 한 줄 요약 ────────────────────────────────────────────────────
    doc.add_heading("1. 요약", level=1)
    _p(doc, f"12개 해석이 요구하는 (재료 × 필수물성) 격자 {tot_cells:,}칸 중 "
            f"{tot_filled:,}칸이 채워졌다 — 전체 {tot_pct:.1f}%.", bold=True, size=11)

    _p(doc, "“몇 %”의 분모를 두 가지로 낸다. 둘은 다른 것을 잰다.", space=2)
    _tbl(doc, ["지표", "정의", "무엇을 재는가"], [
        ["셀 채움률", "채워진 (재료 × 필수물성) 칸 / 전체 칸", "물성 수집의 진척도"],
        ["재료 준비율", "그 해석을 실제로 돌릴 수 있는 재료 / 대상 재료", "해석 가능성"],
    ])
    doc.add_paragraph()
    _p(doc, "셀 채움률이 높아도 재료 준비율은 낮을 수 있다. 재료마다 비어 있는 칸이 다르면 "
            "그렇게 된다 — 낙하가 정확히 그 경우다(셀 53.6% vs 준비 5.0%). "
            "**수집 진척과 해석 가능성을 같은 숫자로 말하면 안 된다.**", bold=True)

    # ── 2. 해석별 커버리지 ───────────────────────────────────────────────
    doc.add_heading("2. 해석별 커버리지", level=1)
    rows = []
    for x in cov:
        top = x["missing"][0] if x["missing"] else ("—", 0)
        label = top[0].replace("택일군: ", "").split(".", 1)[-1]
        rows.append([x["name"], x["n_target"],
                     f"{x['filled']}/{x['cells']}",
                     f"{x['cell_pct']:.1f}%  {bar(x['cell_pct'], 14)}",
                     f"{x['n_ready']}  ({x['ready_pct']:.1f}%)",
                     f"{label[:40]} — {top[1]}종"])
    _tbl(doc, ["해석", "대상", "셀", "셀 채움률", "재료 준비", "가장 큰 공백"], rows, bsz=7.5)
    doc.add_paragraph()

    _p(doc, "읽는 법", bold=True, space=2)
    _tbl(doc, ["구간", "해석", "의미"], [
        ["준비율 50% 이상", "구조·강성 60.8% · 열응력 56.2% · 초탄성 54.3% · 열전달 51.0%",
         "기본 물성만 쓰는 해석. 지금 바로 돌아간다"],
        ["준비율 25~50%", "접착·박리 32.4% · 전기·EMI 27.4% · 결로(투습) 25.7%",
         "부품군을 좁히면 실용 가능"],
        ["준비율 10% 미만", "결로(표면) 7.6% · 폴더블 벤딩 5.4% · 낙하 5.0% · 피로 1.0% · 산소소광 0%",
         "그 해석을 그 해석답게 만드는 물성이 비어 있다"],
    ])
    doc.add_paragraph()
    _p(doc, "낮은 쪽 다섯이 이 카탈로그의 실제 병목이다. 기본 물성(밀도·영률·포아송·k·cp·CTE)은 "
            "57~79%까지 왔는데, 율속·점탄성·젖음성·피로계수·광물리처럼 그 해석의 정체성을 이루는 "
            "물성은 한 자릿수에 머문다.", bold=True)

    # ── 3. 해석별 상세 ───────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("3. 해석별 상세 — 무엇이 막고 있나", level=1)
    for i, x in enumerate(cov, start=1):
        spec = next(a for a in ANALYSES if a[0] == x["name"])
        doc.add_heading(f"3.{i}  {x['name']}", level=2)
        _p(doc, spec[4], space=2)
        _tbl(doc, ["대상 재료", "셀 채움", "재료 준비"],
             [[f"{x['n_target']}종", f"{x['filled']}/{x['cells']}  ({x['cell_pct']:.1f}%)",
               f"{x['n_ready']}종  ({x['ready_pct']:.1f}%)"]])
        if x["missing"]:
            doc.add_paragraph()
            _p(doc, "비어 있는 칸", bold=True, space=2)
            _tbl(doc, ["물성 / 택일군", "미보유 재료 수"],
                 [[k.replace("택일군: ", "").split(".", 1)[-1][:78], f"{v}종"]
                  for k, v in x["missing"]])
        doc.add_paragraph()

    # ── 4. 물성별 보유율 ─────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("4. 물성별 보유율 — 무엇이 채워졌고 무엇이 비었나", level=1)
    _p(doc, f"물성 정의 {n_def}종 전부에 값이 하나 이상 있다(개통률 100%). "
            f"아래는 재료 {n_mat}종 기준 보유율이다.")

    ordered = sorted(props, key=lambda p: -p["n_mat"])
    doc.add_heading("4.1  잘 채워진 것 — 상위 20종", level=2)
    _tbl(doc, ["물성", "도메인", "재료", "보유율", "값"],
         [[p["name"][:34], DOMAIN_KO.get(p["domain"], p["domain"]), f"{p['n_mat']}종",
           f"{p['pct']:.1f}%  {bar(p['pct'], 12)}", f"{p['n_val']}건"]
          for p in ordered[:20]], bsz=7.5)
    doc.add_paragraph()

    doc.add_heading("4.2  비어 있는 것 — 하위 25종", level=2)
    _p(doc, "값이 1~3종에만 있는 물성이다. 대부분 “세트로 있어야 쓸 수 있는” 모델 계수라, "
            "재료 수가 적으면 그 모델을 쓸 수 있는 재료가 그만큼뿐이라는 뜻이다.")
    _tbl(doc, ["물성", "도메인", "재료", "값"],
         [[p["name"][:38], DOMAIN_KO.get(p["domain"], p["domain"]), f"{p['n_mat']}종", f"{p['n_val']}건"]
          for p in ordered[-25:]], bsz=7.5)
    doc.add_paragraph()

    doc.add_heading("4.3  도메인별 분포", level=2)
    dom = defaultdict(lambda: [0, 0, 0])
    for p in props:
        d = dom[p["domain"]]
        d[0] += 1
        d[1] += p["n_val"]
        d[2] += p["n_mat"]
    _tbl(doc, ["도메인", "정의", "물성값", "물성당 평균 재료"],
         [[DOMAIN_KO.get(k, k), f"{v[0]}종", f"{v[1]:,}건", f"{v[2]/v[0]:.0f}종"]
          for k, v in sorted(dom.items(), key=lambda x: -x[1][1])])
    doc.add_paragraph()

    # 카테고리별
    doc.add_heading("4.4  재료 카테고리별 밀도", level=2)
    rows = []
    for cat, n in c.execute("select category,count(*) from material group by category order by count(*) desc"):
        v = c.execute("select count(*) from property_value p join material m on m.id=p.material_id "
                      "where m.category=?", (cat,)).fetchone()[0]
        rows.append([cat, f"{n}종", f"{v:,}건", f"{v/n:.1f}"])
    _tbl(doc, ["카테고리", "재료", "물성값", "재료당 평균"], rows)
    doc.add_paragraph()
    _p(doc, "폼이 재료당 12.8건으로 가장 얇다. 폼은 벤더가 CFD(압축력 편향)와 밀도만 공표하고 "
            "탄성계수·열물성을 싣지 않는 것이 업계 관행이라, 이 얇음은 수집 부족이 아니라 "
            "구조적 부재에 가깝다.")

    # ── 5. 데이터 품질 ───────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("5. 데이터 품질", level=1)
    tiers = dict(c.execute("select quality_tier,count(*) from property_value group by quality_tier"))
    tot = sum(tiers.values())
    _tbl(doc, ["등급", "정의", "건수", "비율"], [
        ["tier 1", "그 재료를 직접 측정해 인쇄한 값(벤더 TDS·논문 실측)",
         f"{tiers.get(1,0):,}", f"{tiers.get(1,0)*100/tot:.1f}%"],
        ["tier 2", "핸드북·표준·공인 DB", f"{tiers.get(2,0):,}", f"{tiers.get(2,0)*100/tot:.1f}%"],
        ["tier 3", "클래스 대표값·2차 인용·디지타이즈", f"{tiers.get(3,0):,}", f"{tiers.get(3,0)*100/tot:.1f}%"],
        ["tier 4", "계산·추정·가정값", f"{tiers.get(4,0):,}", f"{tiers.get(4,0)*100/tot:.1f}%"],
    ])
    doc.add_paragraph()
    n_assum = c.execute("select count(*) from property_value where "
                        "instr(coalesce(conditions,''),'assumption')>0").fetchone()[0]
    _p(doc, f"가정값 {n_assum:,}건은 조건에 assumption 표지가 붙어 UI·카드에 “가정”으로 표시된다. "
            f"대표값 선택에서 실측에 밀리므로, 나중에 실측이 들어오면 자동으로 대체된다.")

    kinds = dict(c.execute("select kind,count(*) from source group by kind order by count(*) desc"))
    _p(doc, "출처 종류", bold=True, space=2)
    _tbl(doc, ["종류", "건수"], [[k, f"{v:,}"] for k, v in
                                sorted(kinds.items(), key=lambda x: -x[1])])
    doc.add_paragraph()

    _p(doc, "정합성 검사 32항목을 통과한 상태다. 주요 항목만 적으면 — 출처 없는 값 0, 단위가 정의와 "
            "다른 값 0, Prony·Darveaux·Anand 등 다항 상수의 항번호 누락 0, 계수·지수 쌍 불일치 0, "
            "Mooney-Rivlin 초기 전단탄성률 음수(Drucker 불안정) 0, LCSR 곡선 가로축 비단조 0.")

    # ── 6. 우선순위 ─────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("6. 무엇을 먼저 채워야 하나", level=1)
    _p(doc, "재료 준비율을 가장 크게 올리는 순서다. 괄호 안은 그 물성 하나가 막고 있는 재료 수.")
    _tbl(doc, ["순위", "물성", "막고 있는 것", "성격"], [
        ["1", "율속(Cowper-Symonds · DIF · LCSR · JC-C)",
         "낙하 준비율 5.0% — 514종이 율속 경로 전무", "수집 여지 있음. 구리·황동은 원출처 미색인"],
        ["2", "비열 · 열전도율",
         "열전달 51% · 결로 두 갈래를 동시에 막는다", "CCL·자성재·커버글라스는 벤더가 아예 안 싣는다"],
        ["3", "점탄성 Prony 세트",
         "벤딩 준비율 5.4% — 175종 전무", "OA에 OCA·무색 PI가 없다. 유료 접근 또는 자체 DMA"],
        ["4", "젖음성(접촉각 · 표면에너지)",
         "결로(표면) 7.6% — 500종 전무", "커버글라스는 문헌 자체가 없다. 세정 이력 명시본만 유효"],
        ["5", "피로계수(Basquin · Coffin-Manson · Darveaux · Morrow)",
         "피로 준비율 1.0%", "솔더·구조금속에 집중하면 소수로 큰 효과"],
        ["6", "광물리(τ₀ · k_q) · 산소 수송(D · S)",
         "산소 소광 0% — 신규 영역", "적색발광 과제. 별도 문서 참조"],
    ], bsz=7.5)
    doc.add_paragraph()

    doc.add_heading("6.1  채울 수 없는 것 — 확인된 부재", level=2)
    _p(doc, "**“안 찾아봤다”와 “찾아봤는데 안 실려 있다”는 다른 상태이고, 후자만 데이터셋 항목이 된다.** "
            "아래는 후자다 — 1차 문서를 끝까지 읽고 그 물성이 인쇄·측정된 적 없음을 확인한 것들이다. "
            "수집을 더 해도 안 나오므로 실측 위탁이나 유료 접근이 아니면 방법이 없다.", bold=True)
    _p(doc, "이 구분이 왜 중요한가 — 커버리지 표의 “0건”은 두 상태를 구분하지 않는다. "
            "어디를 더 파야 하는지 정하려면 이 목록을 먼저 빼고 봐야 한다. "
            "전체 목록은 docs/BLOCKED_SOURCES.md에 있다.")
    _tbl(doc, ["항목", "이유"], [
        ["CCL·프리프레그의 비열", "벤더 6곳 전수 확인. IPC 시험표에 비열 항목 자체가 없다"],
        ["연자성·영구자석의 비열", "METGLAS·VAC·MMPA 데이터시트 15종 확인. 물성 블록에 그 행이 없다"],
        ["커버글라스의 접촉각", "Europe PMC에서 관련 검색이 사실상 0건. 문헌이 존재하지 않는다"],
        ["접착테이프의 열물성·기체투과", "14개 제품군 확인. TDS가 peel·shear·SAFT로 스키마가 고정돼 있다"],
        ["박막봉지(TFE)의 정밀 산소투과도", "요구사양이 상용 계측기 하한 아래라 업계가 산소는 안 재고 수분만 본다"],
        ["박막(파릴렌 등)의 율속 데이터", "두께 1~15 µm라 시험법 자체에 제어된 변형률속도가 없다"],
        ["아크릴 OCA의 산소 확산·용해도", "OCA는 배리어 부품이 아니라 벤더가 측정·공개할 동기가 없다"],
        ["제조사·등급 없는 제네릭 항목 8종", "대조할 원문이 존재하지 않는다. 카탈로그 명명의 문제다"],
        ["CCL·FCCL의 열전도율",
         "TAIFLEX·Chukoh·ITEQ·Nan Ya·TUC 7종의 벤더 발행 문서를 전수 확인. 열전도율 행이 없다. "
         "비열에 이어 열전도율도 같다"],
        ["DELO MONOPOX AC 시리즈의 열전도율",
         "6등급 TDS를 Wayback CDX로 전수 확보해 확인. 열 물성 행 자체가 없다"],
        ["Liu 2013 학위논문의 기판 열전도율",
         "7,174행 전문에 열전도율 문자열 0건. 순수 열역학적 논문이라 측정하지 않았다"],
        ["Sandia 언더필의 열전도율",
         "배합이 일치함을 확인하고도 표가 전부 압축탄성률·Prony·CTE다. 측정된 적이 없다"],
    ])
    doc.add_paragraph()
    _p(doc, "확인 과정에서 걸러낸 함정도 남겨 둔다. 값이 없는 것보다 **잘못된 값이 들어오는 경로**가 "
            "더 위험하기 때문이다.", bold=True, space=2)
    _tbl(doc, ["함정", "실제 정체"], [
        ["TUC 발표자료의 `2 W/mK · 1 W/mK`",
         "같은 슬라이드 **옆 레이아웃 박스**의 절연금속기판(IMS) 계열 값. TU-933·TU-872 SLK 것이 아니다"],
        ["`Thermal Resistance — T260 / T288`",
         "박리까지 걸린 **시간(분)**이다. 길이 차원이 없어 역수를 취해도 열전도율이 안 된다"],
        ["방향이 갈린 행(X/Y/Z)", "열전도율이 아니라 **CTE**다. 일곱 종 어디에도 열전도율 방향 분리는 없다"],
        ["Chukoh 사이트의 일반 PTFE 수지 k", "수지 값이지 유리섬유 충전 CCL 값이 아니다"],
        ["`DELO MONOPOX 1.7 W/(m·K)`",
         "**TC2270**(열전도 등급)의 값. 대상인 **AC**(이방도전) 시리즈와 다른 제품 라인이다"],
        ["CEL-9220HF10 0.9 · CEL-8240HF10GK 1.0 W/mk",
         "실측이 확인된 값이지만 **CEL-9240HF와 다른 등급**이다. 대체하면 안 된다"],
    ], bsz=7.5)
    doc.add_paragraph()

    # ── 7. 지표 설계 ────────────────────────────────────────────────────
    doc.add_heading("7. 이 통계를 읽을 때 주의할 것", level=1)
    _p(doc, "이 프로젝트에서 지표가 준비도를 부풀린 사고가 세 번 있었다. 그때마다 고쳤고, "
            "지금 숫자는 그 수정을 반영한 것이다.")
    _tbl(doc, ["사고", "내용", "고친 방식"], [
        ["낙하 지표", "GF 강화 등급은 항복점을 재지 않는 것이 정상인데 결측으로 세어 146종이 "
                    "부당하게 미달로 잡혔다", "‘항복 OR (인장 AND 연신)’ 택일로 바꿨다. 22% → 74%"],
        ["벤딩 지표 (1차)", "Prony 키 보유로 세면 부풀려진다. 항이 있어도 τ와 항번호가 짝을 "
                        "이루지 않으면 곡선이 안 된다", "실제 생성기를 돌려 세도록 바꿨다"],
        ["벤딩 지표 (2차)", "세트가 일관돼도 밀도·체적탄성률이 없으면 카드가 안 나온다",
         "Prony 보유 / 세트 일관 / 덱 생성 세 단계로 나눠 표시한다"],
    ], bsz=7.5)
    doc.add_paragraph()
    _p(doc, "교훈은 하나다 — **물성을 보유했다는 것과 해석에 쓸 수 있다는 것은 다르다.** "
            "이 문서의 ‘재료 준비율’은 후자를 재려 한 것이고, ‘셀 채움률’은 전자를 잰다.", bold=True)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"저장: {OUT}")
    print(f"  해석 {len(cov)}개 · 물성 {len(props)}종 · 전체 셀 채움 {tot_pct:.1f}%")


if __name__ == "__main__":
    main()
