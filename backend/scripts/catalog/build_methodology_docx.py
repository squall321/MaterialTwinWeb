#!/usr/bin/env python
# 물성 수집 방법론(빈칸 탐지 → 채움 → 검증 → 신뢰도 고도화)을 docx 한 부로 묶는다.
# 수치는 전부 라이브 DB에서 계산한다 — 하드코딩하면 파동마다 문서가 어긋난다.
# 사용: .venv/bin/python scripts/catalog/build_methodology_docx.py
from __future__ import annotations

import sqlite3
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

import coverage_report as CR  # noqa: E402
from coverage_report import ANALYSES, compute  # noqa: E402

OUT = Path("/home/koopark/claude/MaterialTwinWeb/MaterialTwin_물성수집_방법론.docx")
DB = "/home/koopark/claude/HEAXHub/var/app_data/materialtwin_web/materialtwin.db"

# 문서화된 method×tier 허용 조합(PROPERTY_DATA_HISTORY "정한 규칙" 2항).
# 이 밖의 조합이 나오면 어딘가 어긋난 것이다 — 문서가 매번 세어서 드러낸다.
METHOD_TIER_OK = {("measured", 1), ("measured", 3), ("handbook", 2), ("handbook", 3),
                  ("computed", 3), ("computed", 4), ("estimated", 4)}

# 레버리지 항목의 읽을 수 있는 이름과 "왜 어려운가".
# 택일군 라벨은 taxonomy 키를 이어붙인 것이라 그대로 실으면 표에서 잘린다.
# (키에 포함된 문자열, 한글 라벨, 사유)
LEVER_MAP = [
    ("cowper_symonds", "율속 (Cowper-Symonds · DIF · σy(ε̇) · JC-C)",
     "OA에 존재하는 Cowper-Symonds 논문이 45건 수준이고 전부 금속 아니면 벌크 폴리머다. "
     "복합재·세라믹·유리의 CS 쌍은 OA에 하나도 없다"),
    ("yield_strength |", "소성 (항복 OR 인장+연신)",
     "적층판 18종은 파단연신율을 어느 벤더도 발표하지 않고 문헌에도 없다"),
    ("fatigue_strength_coefficient", "피로계수 (Basquin · Coffin-Manson · Darveaux · Morrow)",
     "구리 계열 9종은 공표된 상관식에 분기가 아예 없다. FD&E 원시 데이터는 파일 자신이 "
     "“digitized from plots may contain errors”라고 밝힌다 — 우리가 피팅하면 역산이다"),
    ("water_vapor_transmission", "투습 (WVTR · 투습계수 · 확산 · 흡습)",
     "저분자 유기 고체는 자립 필름을 못 만들어 시험 자체가 성립하지 않는다"),
    ("peel_strength", "접착 (peel · lap shear · die shear)",
     "동박 peel은 동박 단독이 아니라 수지와의 쌍 물성이라 벤더가 안 싣고 IPC-4562A에도 표가 없다"),
    ("dielectric_constant", "전기 (유전율 · 체적저항률 · 손실계수)",
     "도체는 유전율이 정의되지 않고, 저밀도 폼은 상대밀도가 지배하는데 벤더가 발표하지 않는다"),
    ("prony_relaxation_time", "점탄성 Prony 세트",
     "상용 PI 필름·점착테이프의 점탄성은 Prony로 발표되지 않는다(ETAG 002 크리프로 규격화). "
     "사내 DMA 없이는 안 열린다"),
    ("gas_permeability_o2", "산소 (O₂ 투과계수 · 용해도)", "OCA는 배리어 부품이 아니라 벤더가 측정할 동기가 없다"),
    ("contact_angle_water", "젖음 (접촉각 · 표면에너지)", "커버글래스는 bare 값을 상한으로만 인쇄한다"),
    ("hyperelastic_coefficient", "초탄성 계수", "압축성 폼 초탄성 피팅이 공개된 사례가 없다"),
    ("youngs_modulus", "영률", "남은 2종은 개별 확인 대상이다"),
    ("thermal.conductivity", "열전도율", "CCL·FCCL은 벤더 전수 확인으로 부재가 확정됐다"),
    ("thermal.expansion_linear", "선팽창계수(CTE)", "남은 1종은 개별 확인 대상이다"),
    ("thermal.specific_heat", "비열", "CCL·테이프·연자성재는 어느 벤더도 발표하지 않는다"),
]


def lever_label(key: str) -> tuple[str, str]:
    """레버리지 키 → (한글 라벨, 왜 어려운가). 매칭이 없으면 키를 그대로 쓴다."""
    for frag, label, why in LEVER_MAP:
        if frag in key:
            return label, why
    return key.replace("택일군: ", "")[:44], "—"


# ── 서식 헬퍼 (build_coverage_docx.py와 동일 규격) ──────────────────────────
def _sty(doc):
    n = doc.styles["Normal"]
    n.font.name = "맑은 고딕"
    n.font.size = Pt(9.5)
    n.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    for i, sz in ((1, 15), (2, 12.5), (3, 11)):
        h = doc.styles[f"Heading {i}"]
        h.font.name = "맑은 고딕"
        h.font.size = Pt(sz)
        h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        h.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def _tbl(doc, headers, rows, hsz=8.5, bsz=8):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(str(h))
        r.bold = True
        r.font.size = Pt(hsz)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(bsz)
    return t


def _p(doc, text, bold=False, size=9.5, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    return p


def bar(pct, width=20):
    """텍스트 막대 — 표에서 한눈에 비교되게."""
    n = int(round(pct / 100 * width))
    return "█" * n + "·" * (width - n)


# ── 정직한 격자 통계 ────────────────────────────────────────────────────────
def honest_grid():
    """부재 칸에 앉은 tier4를 분자에서도 빼고 다시 센다.

    `coverage_report`의 유효채움은 부재 칸을 분모에서 빼면서 분자에는 남겨 둬,
    해석에 따라 100%를 넘는다. 남은 일의 크기를 말하려면 양쪽을 같은 집합에서 세야 한다.
    """
    c, mat, own, meas = CR.load()
    CR._build_active(mat, own, meas)
    rows, lever, short_all = [], Counter(), Counter()
    lever_mats = defaultdict(set)
    tot = dict(cells=0, filled=0, unfill=0, unfill_filled=0)
    for name, must, anyof, filt, _desc in ANALYSES:
        tgt = CR.scope(mat, filt)
        cells = filled = unfill = unfill_filled = 0
        gap_by_mat = Counter()
        for m in tgt:
            ks = own[m]
            for k in must:
                cells += 1
                u = CR._cell_unfillable(mat[m][0], (k,), mat[m][1])
                unfill += u
                if k in ks:
                    filled += 1
                    unfill_filled += u
                elif not u:
                    gap_by_mat[m] += 1
                    lever[k] += 1
                    lever_mats[k].add(m)
            for grp in anyof:
                cells += 1
                u = CR._cell_unfillable(mat[m][0], grp, mat[m][1])
                unfill += u
                if CR.grp_ok(ks, grp):
                    filled += 1
                    unfill_filled += u
                elif not u:
                    gap_by_mat[m] += 1
                    lbl = "택일군: " + " | ".join(
                        ("+".join(x.split(".", 1)[1] for x in alt) if isinstance(alt, tuple)
                         else alt.split(".", 1)[1]) for alt in grp)
                    lever[lbl] += 1
                    lever_mats[lbl].add(m)
        for m, n in gap_by_mat.items():
            short_all[m] += n
        d = Counter(gap_by_mat.values())
        rows.append({"name": name, "n": len(tgt), "cells": cells, "filled": filled,
                     "unfill": unfill, "unfill_filled": unfill_filled,
                     "eff": cells - unfill, "eff_filled": filled - unfill_filled,
                     "gap": (cells - unfill) - (filled - unfill_filled),
                     "left_mat": len(gap_by_mat), "s1": d.get(1, 0), "s2": d.get(2, 0),
                     "s3": sum(v for k, v in d.items() if k >= 3)})
        for k in tot:
            tot[k] += rows[-1][k] if k in rows[-1] else 0
        tot["cells"] = sum(r["cells"] for r in rows)
        tot["filled"] = sum(r["filled"] for r in rows)
        tot["unfill"] = sum(r["unfill"] for r in rows)
        tot["unfill_filled"] = sum(r["unfill_filled"] for r in rows)
    return rows, tot, lever, lever_mats, short_all


def main():
    c, mat, own, cov = compute()
    hrows, htot, lever, lever_mats, short_all = honest_grid()
    by_name = {r["name"]: r for r in hrows}

    q = lambda s, *a: c.execute(s, *a).fetchall()
    n_mat, n_val, n_src, n_def = c.execute(
        "select (select count(*) from material),(select count(*) from property_value),"
        "(select count(*) from source),(select count(*) from property_definition)").fetchone()
    tiers = dict(c.execute("select quality_tier,count(*) from property_value group by 1"))
    n_assum = c.execute("select count(*) from property_value where replace(coalesce("
                        "conditions,''),' ','') like '%\"assumption\":true%'").fetchone()[0]
    mt = q("select method,quality_tier,count(*) from property_value group by 1,2")
    odd = [(m, t, n) for m, t, n in mt if (m, t) not in METHOD_TIER_OK]

    tot_cells = sum(x["cells"] for x in cov)
    tot_filled = sum(x["filled"] for x in cov)
    tot_meas = sum(x["meas_filled"] for x in cov)
    true_gap = htot["cells"] - htot["unfill"] - (htot["filled"] - htot["unfill_filled"])
    dead = htot["unfill"] - htot["unfill_filled"]

    doc = Document()
    _sty(doc)

    # ── 표지 ─────────────────────────────────────────────────────────────
    h = doc.add_heading("MaterialTwin 물성 수집 방법론", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("빈칸을 어떻게 찾아내고 · 채우고 · 검증하고 · 더 믿을 만한 출처로 갈아 끼우는가")
    r.font.size = Pt(13)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{date.today().isoformat()} · 재료 {n_mat:,}종 · 물성값 {n_val:,}건 · "
                  f"출처 {n_src:,}건 · 물성 정의 {n_def}종")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()
    _p(doc, "이 문서는 규칙집이 아니라 아홉 번의 파동을 돌리며 실제로 깨지고 고쳐진 "
            "프로세스의 기록이다. 수치는 전부 라이브 DB에서 계산했다. "
            "경위의 1차 사료는 docs/PROPERTY_DATA_HISTORY.md(72개 장)이고, "
            "수집 지시서 정본은 docs/COLLECTION_BRIEF.md다.", space=2)
    _p(doc, "읽는 사람은 셋을 가정한다 — 해석 입력을 받아 쓰는 솔버팀, 값을 채우는 물성팀, "
            "“지금 몇 %이고 남은 일이 얼마인가”를 묻는 경영진.")

    # ── 1. 문제 정의 ─────────────────────────────────────────────────────
    doc.add_heading("1. 문제 정의 — 왜 물성 카탈로그인가", level=1)
    _p(doc, "해석 하나를 돌리려면 그 해석이 요구하는 물성이 전부 있어야 한다. "
            "아홉 개가 있고 하나가 없으면 결과는 90%가 아니라 0%다. 덱이 안 만들어진다.", bold=True)
    _p(doc, "이 성질이 모든 판단을 지배한다. 물성값 총량을 늘리는 것과 해석 가능한 재료를 "
            "늘리는 것은 다른 일이다. 재료마다 마지막 한 칸이 다르면 총량이 아무리 늘어도 "
            "돌아가는 해석은 늘지 않는다. 8차 파동에서 방침이 바뀐 이유도 이것이다 — "
            "덱이 도는 것과 안 도는 것의 차이가, 그 값이 tier1이냐 tier3이냐의 차이보다 크다.")

    doc.add_heading("1.1  격자 — 무엇을 세고 있나", level=2)
    _p(doc, f"카탈로그는 재료 {n_mat}종 × 물성 정의 {n_def}종이다. 그러나 이 8만여 칸을 다 세지 "
            "않는다. 알루미나에 투습도를 요구하거나 벌크 금속에 층두께를 요구하면 없는 결함을 "
            "만들어내는 지표가 된다(실제로 그렇게 세다가 벤딩을 8%, 결로를 4%로 잘못 셌다). "
            f"대신 해석 {len(cov)}종이 실제로 요구하는 (재료 × 필수물성) 격자 "
            f"{tot_cells:,}칸을 센다.", bold=True)
    _tbl(doc, ["해석", "대상", "필수물성", "택일군(하나만 있어도 됨)", "칸"],
         [[x["name"], f"{x['n_target']}종",
           ", ".join(k.split(".", 1)[1] for k in spec[1]) or "—",
           " / ".join(
               " · ".join(("+".join(y.split(".", 1)[1] for y in alt) if isinstance(alt, tuple)
                           else alt.split(".", 1)[1]) for alt in g) for g in spec[2])[:56] or "—",
           f"{x['cells']:,}"]
          for x, spec in ((x, next(a for a in ANALYSES if a[0] == x["name"])) for x in cov)],
         bsz=7)
    doc.add_paragraph()
    _p(doc, "택일군이 있는 이유 — 소성·파괴 입력은 항복강도로도 되고 인장강도+파단연신율로도 "
            "된다. GF 강화 등급은 항복점을 애초에 재지 않는 것이 정상인데 이걸 결측으로 세는 "
            "바람에 146종이 부당하게 미달로 잡혀 낙하 준비율이 22%로 찍힌 적이 있다. "
            "고치니 74%였다.")

    # ── 2. tier 체계 ─────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("2. tier 체계 — 등급이 무엇을 뜻하고 왜 tier4를 허용하는가", level=1)
    _p(doc, "tier의 기준은 “문서의 종류”가 아니라 “값의 성격”이다. 벤더 TDS든 논문이든 "
            "상관없다.", bold=True)
    _tbl(doc, ["tier", "뜻", "건수", "비율"], [
        ["1", "그 제품·시편에 대해 문서에 인쇄된 실제 시험값(벤더 TDS·그 재료를 시험한 논문)",
         f"{tiers.get(1,0):,}", f"{tiers.get(1,0)*100/n_val:.1f}%"],
        ["2", "핸드북·표준 최소보증치·공인 DB(MMPDS·ASM·NIST·CDA·ISO/ASTM)",
         f"{tiers.get(2,0):,}", f"{tiers.get(2,0)*100/n_val:.1f}%"],
        ["3", "계열·클래스 대표값 / 2차 인용 / 등급 불일치 대체값 / 그래프 좌표 추출",
         f"{tiers.get(3,0):,}", f"{tiers.get(3,0)*100/n_val:.1f}%"],
        ["4", "계산·유도·추정·가정", f"{tiers.get(4,0):,}", f"{tiers.get(4,0)*100/n_val:.1f}%"],
    ])
    doc.add_paragraph()

    doc.add_heading("2.1  tier4를 허용하되 반드시 표시하는 이유", level=2)
    _p(doc, "7차까지는 “빈 칸이 틀린 값보다 낫다”를 가장 강하게 해석했다. 8차에서 방침이 "
            "바뀌었다 — 계열 대표값(tier3)과 근거 있는 가정(tier4)까지 허용해 일단 칸을 닫는다. "
            "다만 바뀐 것은 어느 tier까지 허용하느냐이지 근거 없이 지어내도 되느냐가 아니다.")
    _p(doc, "그래서 tier4에는 두 개의 구조적 표지를 필수로 박는다 — "
            "conditions.assumption = true 와 method = \"estimated\".", bold=True)
    _tbl(doc, ["장치", "어디서 작동하나", "효과"], [
        ["대표값 선택(_rep_rank)", "조회 · 비교 · DYNA 카드",
         "tier가 낮은 값이 이긴다. 가정값은 항상 실측에 밀린다"],
        ["UI 배지(tierBadge)", "화면",
         "tier4 + assumption이면 “계산/등가”가 아니라 “가정”(danger 색)으로 표시"],
        ["무결성 검사 2항목", "배포 전",
         "가정값인데 tier4·estimated 아님 / 가정값인데 근거 출처 없음"],
    ])
    doc.add_paragraph()
    _p(doc, "핵심은 자동 대체다. 나중에 실측이 들어오면 그 순간 대표값이 바뀐다 — 사람이 손으로 "
            "지울 필요가 없다. 실제로 작동한 것을 확인한 사례가 둘 있다.", bold=True)
    _tbl(doc, ["사례", "무엇이 바뀌었나"], [
        ["배터리 분리막 포아송비", "tier4 추정 0.4 → tier2 실측 0.17로 대표값이 갈아 끼워졌다"],
        ["Al7050-T7451 피로계수",
         "tier4 추정(Meggiolaro-Castro 1.9×Su = 996 MPa) → tier3 실측 계열대체"
         "(AA7075-T6 FD&E 피팅 876 MPa)"],
    ])
    doc.add_paragraph()
    _p(doc, "반대 방향의 함정도 있다. 가정 위에 가정을 쌓으면 소용없다. 포아송비 가정값을 만들 때 "
            "처음엔 전체 값으로 클래스 중앙값을 냈는데, 폼 4건이 전부 0.300이고 편차가 0이라 "
            "보니 이미 tier4 가정값이었다. 앵커로 쓸 행은 값과 tier만 보지 말고 method와 notes를 "
            "반드시 읽어야 한다.")

    # ── 3. 넘지 않는 선 ─────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("3. 절대 넘지 않는 선 — 그리고 그 규칙이 막아 온 사고", level=1)
    _p(doc, "규칙은 추상이 아니다. 전부 실제로 사고가 나서 생겼다.", bold=True)
    _tbl(doc, ["규칙", "그 규칙이 없었으면 난 사고"], [
        ["1. 지어내지 않는다. 빈 칸이 틀린 값보다 낫다",
         "라미네이트 밀도 16종을 비워 뒀다 — Isola 370HR SDS가 "
         "‘SPECIFIC GRAVITY: Not Available for product’라고 명기한다. "
         "Kaneka Graphinity의 1200 W/mK는 무역기사에만 뜨는 숫자라 채택하지 않았다. "
         "8차는 항복강도 258종 중 69종만 닫았고 그게 정답이었다 — 171종은 항복강도가 "
         "물리적으로 존재하지 않는다"],
        ["2. 역산 금지",
         "VHB 영률 0.894 MPa는 E=2G(1+ν) 역산값인데 notes에는 “Satas 핸드북 계열값”이라 "
         "적혀 있었다. 역산값 하나가 13종으로 번지고 근거 표기까지 틀린 채로 남았다. "
         "AZoM/Granta의 폴리머 G는 5계열에서 E/(2(1+ν))와 유효숫자까지 일치하고, "
         "MakeItFrom의 LDPE는 E·G 조합이 ν=−0.29를 만든다"],
        ["3. 그래프에서 값을 읽지 않는다",
         "아크릴 PSA 열전도율 앵커 0.176은 인용 논문 전문에 0회 등장한다(순수 PSA 값은 "
         "Figure에만 있다). SUS430 율속은 원문이 “그림 4에서 읽었다”여서 전량 폐기. "
         "다만 pdftotext가 빈 표를 주면 그것은 이미지 표이므로 렌더해 읽는 것은 규칙 3에 "
         "걸리지 않는다"],
        ["4. 재인쇄본·2차 인용을 믿지 않는다",
         "Polymers 12(10) 2327 Table 4는 캡션이 2차 인용을 밝히고 같은 표의 PC 밀도 1.44가 "
         "실제 1.20과 어긋난다(19건에 검산 실패를 박았다). Kwatra가 Ferguson 값을 8.18로 "
         "오기(원문 8.48). PV 적층 물성표는 세 문헌에 바이트 단위로 복사돼 있었다 — 같은 "
         "숫자 조합이 여러 논문에 나오면 독립 확인이 아니라 재인쇄다"],
        ["5. 조건 없는 값은 값이 아니다",
         "3M VHB 4910 TDS의 Shear Modulus 6e5 Pa는 시험규격도 율속도 없는데 같은 재료의 "
         "율속 분해 문헌값보다 12~35배 크다. tier1로 두면 조건 없는 값이 대표로 뽑혀 조건 "
         "붙은 값을 밀어낸다. 파장 없는 굴절률 75건도 같은 이유로 tier4로 내렸다"],
        ["6. 상위 tier 값과 자릿수가 어긋나면 넣지 않는다",
         "Foam PU의 5항 Prony는 수집이 흠잡을 데 없었으나 그 재료의 tier1 밀도가 240 kg/m³인데 "
         "논문 시료는 65다. 전해동박 항복 467~600 MPa는 그 재료의 tier1 인장강도 147~276 MPa의 "
         "2배다. 등급 불일치는 tier를 낮춰 쓸 수 있지만 이건 카드를 물리적으로 앞뒤 안 맞게 "
         "만든다"],
        ["7. 인접 열 오독을 조심한다",
         "누적 20건 넘게 잡혔다 — 열전도율을 비열로, (BH)max를 굽힘강도로, Chukoh ‘비중 2.15’가 "
         "실은 유전율, Parylene HT 인장이 옆의 Acrylic 열. 공통점은 전부 “그럴듯한” 값이었다는 "
         "것이다. 범위 검사로는 안 걸린다"],
    ], bsz=7.5)
    doc.add_paragraph()
    _p(doc, "9차에서 규칙 2에 한 칸을 열었다(브리프 0-B절). 문헌이 “실측이 없을 때 이렇게 "
            "추정하라”고 공표한 상관식은 쓴다(Universal Slopes · Meggiolaro-Castro Medians · "
            "Gent 등). 금지는 우리가 손으로 세운 항등식이다. 둘의 차이는 누가 그 관계를 "
            "보증하느냐다 — 공표된 상관식은 저자가 적용 범위와 오차까지 밝히고, 손으로 세운 "
            "항등식은 아무도 보증하지 않는다.", bold=True)

    # ── 4. 파동 사이클 ───────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("4. 파동(wave) 사이클 — 한 바퀴에 무엇을 하고 무엇을 검사하나", level=1)
    _p(doc, "한 파동은 대략 하루~이틀, 병렬 배치 6~11갈래다. 순서는 고정이다.")
    _tbl(doc, ["#", "단계", "하는 일", "무엇을 검사하나", "도구"], [
        ["1", "빈칸 분석", "격자에서 비어 있는 칸을 뽑는다",
         "구조적 부재를 뺀 ‘채울 수 있는 칸’인가", "coverage_report.py"],
        ["2", "레버리지 계산", "무엇을 채우면 몇 칸이 닫히나",
         "칸 수가 아니라 칸/재료로 잰다", "—"],
        ["3", "타깃 파일 생성", "배치별 재료 목록",
         "무엇이 없어서 이 재료가 뽑혔는지를 함께 적는다", "—"],
        ["4", "병렬 배치 투입", "브리프 + 타깃 파일 배포",
         "배치마다 scratchpad/<영역>/ 분리", "COLLECTION_BRIEF.md"],
        ["5", "chunk 중간 저장", "재료 10~12종마다 파일을 쓴다",
         "프로세스가 죽어도 남는가", "—"],
        ["6", "JSON 파싱 검증", "json.load()", "파싱 실패 0", "—"],
        ["7", "인제스트 드라이런", "검증만 하고 보고", "6단계 검증(아래)", "ingest_agent_json.py"],
        ["8", "적용", "--apply", "신규 재료·거부 사유 리뷰", "ingest_agent_json.py"],
        ["9", "tier 정규화", "값의 성격으로 등급 재판정",
         "자동 규칙 vs MANUAL/FROZEN 예외", "normalize_tiers.py"],
        ["10", "무결성 검사", "32항목 + 주의 1항목", "0이 아니면 결함", "integrity_check.py"],
        ["11", "통계·문서 재생성", "커버리지·xlsx·docx",
         "하드코딩 없이 전부 계산값인가", "build_* 스크립트"],
        ["12", "이력 기록 → 배포", "무엇이 틀렸는지까지 남긴다", "—", "PROPERTY_DATA_HISTORY.md"],
    ], bsz=7.5)
    doc.add_paragraph()
    _p(doc, "5단계가 왜 규칙이 됐나 — 8차 첫 투입분 6개가 결과를 마지막에 한 번만 쓰도록 돼 "
            "있었고, 프로세스 종료로 산출물이 0이었다. 여섯 배치가 통째로 날아갔다. "
            "장시간 배치는 중간 저장을 설계에 넣어야 한다.", bold=True)

    doc.add_heading("4.1  인제스트가 거부하는 것", level=2)
    _p(doc, "에이전트는 DB에 직접 쓰지 않는다. SQLite 락 경합을 피하고 등록 전 검토 지점을 "
            "만들기 위해 JSON으로 내놓게 하고, 메인이 검증 후 직렬 등록한다.")
    _tbl(doc, ["#", "검증", "무엇을 막나"], [
        ["①", "property_key가 taxonomy에 있는가",
         "그럴싸하지만 없는 키를 지어내는 것(poissons_ratio, expansion_coefficient …)"],
        ["②", "단위가 정의 단위와 일치하는가",
         "곱 순서만 다른 표기는 같게 보되 기호가 다르면 거부(mN/m vs J/m²)"],
        ["③", "출처(url/doi/title)가 있는가",
         "없으면 항목 전체 건너뜀. 스키마 설계 때부터 강제라 “출처 없음”이 끝까지 0"],
        ["④", "물리적 범위 안인가",
         "40여 키에 범위. 신규 키는 범위 자체가 단위 검산이다(CS를 MPa 원값으로 넣으면 걸린다)"],
        ["⑤", "비율 물성이 1을 넘지 않는가", "연신율·흡습률 퍼센트 오입력(100배 오류)"],
        ["⑥", "기존 tier1 실측을 덮지 않는가", "단, 조건이 다르면 별개 측정이므로 통과"],
        ["+", "조건 충돌",
         "조건 서명이 같은데 값이 다르면 덮어쓰기 대신 거부하고 사유 보고. "
         "3M 112P05 박리력이 112P02와 조건이 같아 실제로 조용히 사라져 있었다"],
    ], bsz=7.5)
    doc.add_paragraph()

    doc.add_heading("4.2  무결성 검사 32항목", level=2)
    _p(doc, "배포 전 반드시 0으로 통과시킨다.")
    _tbl(doc, ["검사군", "항목 예"], [
        ["참조 무결성", "출처 없는 값 · 끊어진 출처 참조 · 정의 없는 물성키 · 고아 시편"],
        ["어휘·타입", "tier 범위 · method 어휘 · 출처 kind · category · 조건이 dict 아님"],
        ["단위", "단위가 정의와 다름(곱 순서 정규화 후 비교)"],
        ["조건 필수", "초탄성에 model · 광학값에 파장 · 확산계수에 온도 · 노화 유지율 조건 넷"],
        ["조건 금지", "온도가 값인 물성에 온도 조건(융점·Tg·HDT — 대표값 선택이 무너진다)"],
        ["모델 세트", "Prony 항번호 · Basquin/Coffin-Manson/Morrow 쌍 · Darveaux 홀수 · 항번호 없는 다항 상수"],
        ["물리", "Mooney-Rivlin G0 음수(Drucker 불안정) · 비금속인데 비열이 금속급 · 파장이 가시광 밖"],
        ["가정값 규율", "가정값인데 tier4·estimated 아님 · 가정값인데 근거 출처 없음"],
        ["소비 로직", "LCSR 곡선 가로축 비단조 — SQL로는 못 잡아 실제 생성기를 돌려 확인"],
        ["주의(결함 아님)", "열적 이방성이 입증됐는데 방향 없는 열전도율"],
    ], bsz=7.5)
    doc.add_paragraph()
    _p(doc, "마지막 항목이 중요하다. 방향을 지어내서 검사를 통과시키는 것도, 아는 위험을 안 "
            "알리는 것도 안 된다. 그래서 결함으로 세지 않고 경고로 남긴다.")

    # ── 5. 우선순위 ─────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("5. 우선순위를 어떻게 정하는가", level=1)
    doc.add_heading("5.1  칸 수가 아니라 칸/재료로 잰다", level=2)
    _p(doc, "8차 착수 시점에 빈칸이 5,150개였다. 물성별로 세면 밀도가 893칸으로 압도적 1위다. "
            "그런데 그 893칸의 원인은 재료 134종뿐이었다 — 재료 하나당 6.7칸이다. "
            "밀도는 거의 모든 해석의 필수 물성이라 한 번 채우면 여러 칸이 동시에 닫힌다. "
            "134종에 배치 하나를 붙였더니 전체 셀 채움이 72.2% → 82.3%로 10%p 뛰었다.", bold=True)

    doc.add_heading("5.2  그런데 지금은 그 척도가 무력하다", level=2)
    _p(doc, "같은 계산을 오늘 다시 돌리면 결과가 완전히 뒤집힌다.")
    _tbl(doc, ["물성 / 택일군", "닫히는 칸", "재료 수", "칸/재료"],
         [[lever_label(k)[0], n, len(lever_mats[k]), f"{n/len(lever_mats[k]):.2f}"]
          for k, n in lever.most_common(20)], bsz=7.5)
    doc.add_paragraph()
    _best = max(lever, key=lambda k: lever[k] / len(lever_mats[k]))
    _p(doc, f"칸/재료 최고가 {_best.replace('택일군: ','')[:40]}의 "
            f"{lever[_best]/len(lever_mats[_best]):.2f}인데 대상이 {len(lever_mats[_best])}종뿐이다. "
            "8차의 밀도(6.7칸 × 134종) 같은 덩어리는 더 이상 없다. "
            "기본 물성이 닫히면 이 척도는 스스로 수명을 다한다.", bold=True)

    doc.add_heading("5.3  세 번째 척도 — “1칸 부족 재료”를 센다", level=2)
    _p(doc, "레버리지가 평평해진 지금 실제로 우선순위를 정하는 것은 한 칸만 채우면 해석이 "
            "열리는 재료의 수다.")
    _d = Counter(short_all.values())
    _tbl(doc, ["남은 빈칸", "재료 수"],
         [[f"{k}칸", f"{_d[k]}종"] for k in sorted(_d)] +
         [["빈칸 있는 재료 합계", f"{len(short_all)}종 / {n_mat}종"]])
    doc.add_paragraph()
    _tbl(doc, ["해석", "남은 재료", "1칸 부족", "2칸", "3칸+"],
         [[r["name"], f"{r['left_mat']}종", f"{r['s1']}종", f"{r['s2']}종", f"{r['s3']}종"]
          for r in sorted(hrows, key=lambda x: -x["left_mat"])])
    doc.add_paragraph()
    _only1 = [r["name"] for r in hrows if r["left_mat"] and r["s1"] == r["left_mat"]]
    _p(doc, f"해석 {len(hrows)}개 중 {len(_only1)}개는 남은 재료가 전부 “딱 한 칸” 모자란다. "
            "그 한 칸이 무엇인지도 해석마다 하나로 정해져 있다. 물성 하나를 뚫으면 해석 하나가 "
            "통째로 100%가 되는 자리가 그만큼 열려 있다는 뜻이다.", bold=True)
    _multi = [r for r in hrows if r["s2"] + r["s3"] > 0]
    if _multi:
        _p(doc, "예외는 " + " · ".join(
            f"{r['name']}({r['s2']+r['s3']}종이 2칸 이상)" for r in _multi) +
            "이다. 여기는 한 배치로 안 닫히는 다층 병목이다.")
    _p(doc, "우선순위를 잘못 잡았던 실제 사례도 남긴다. 낙하가 22%로 꼴찌일 때 “항복강도가 "
            "병목”이라 판단해 에이전트 하나를 통째로 썼는데, 결과는 “채울 수 없는 항목”이었다. "
            "지표를 고치니 낙하는 74%였고 진짜 병목은 율속이었다. 지표가 틀리면 우선순위도 "
            "틀린다.", bold=True)

    # ── 6. 지표 넷 ──────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("6. 지표 네 가지와 각각이 감추는 것", level=1)
    _tbl(doc, ["지표", "정의", "재는 것", "감추는 것"], [
        ["셀 채움", "채워진 칸 / 전체 칸", "수집 진척도",
         "값의 품질. tier4 가정값도 채운 것으로 센다"],
        ["실측기반", "tier4를 뺀 칸 / 전체 칸", "근거의 두께",
         "재료마다 다른 칸이 비었다는 사실"],
        ["재료 준비", "그 해석을 돌릴 수 있는 재료 / 대상 재료", "해석 가능성",
         "그 카드가 가정값으로 서 있는지"],
        ["적용 대비 준비", "준비된 재료 / 원리적으로 적용 가능한 재료", "도달 가능한 천장 대비",
         "분모를 줄여 올린 수치라는 점"],
    ])
    doc.add_paragraph()
    _tbl(doc, ["해석", "대상", "셀 채움", "실측기반", "재료 준비", "적용 대비", "실측 기준 준비"],
         [[x["name"], x["n_target"], f"{x['cell_pct']:.1f}%  {bar(x['cell_pct'],8)}",
           f"{x['meas_pct']:.1f}%", f"{x['ready_pct']:.1f}%",
           f"{x['ready_app_pct']:.1f}%", f"{x['ready_meas_pct']:.1f}%"] for x in cov] +
         [["전체", "", f"{tot_filled*100/tot_cells:.1f}%", f"{tot_meas*100/tot_cells:.1f}%",
           "", "", ""]], bsz=7.5)
    doc.add_paragraph()

    doc.add_heading("6.1  지표 하나만 보면 속는 자리", level=2)
    _st = next(x for x in cov if x["name"] == "구조·강성")
    _th = next(x for x in cov if x["name"] == "열전달·방열")
    _wo = max(cov, key=lambda x: x["cell_pct"] - x["ready_pct"])
    _p(doc, f"(가) 셀 채움만 보면 “거의 다 됐다”로 읽힌다. 구조·강성은 셀 {_st['cell_pct']:.1f}%다. "
            f"그런데 실측기반은 {_st['meas_pct']:.1f}%이고, 실측만으로 카드가 나오는 재료는 "
            f"{_st['ready_meas_pct']:.1f}%({_st['n_ready_meas']}종)뿐이다. 전체로는 "
            f"{tot_filled-tot_meas:,}칸({(tot_filled-tot_meas)*100/tot_cells:.1f}%p)이 "
            "가정값이 만든 착시다.", bold=True)
    _p(doc, f"(나) 재료 준비만 보면 “실측으로 서 있는 줄” 안다. 열전달·방열은 준비율 "
            f"{_th['ready_pct']:.1f}%({_th['n_ready']}종)인데 실측 기준으로는 "
            f"{_th['ready_meas_pct']:.1f}%({_th['n_ready_meas']}종)다. "
            f"덱은 {_th['n_ready']}종이 돌지만 그중 {_th['n_ready']-_th['n_ready_meas']}종은 "
            "가정값을 딛고 서 있다.")
    _p(doc, f"(다) 셀 채움과 재료 준비가 크게 갈리는 자리가 진짜 병목이다. "
            f"{_wo['name']}은 셀 {_wo['cell_pct']:.1f}%인데 준비 {_wo['ready_pct']:.1f}%다. "
            "재료마다 비어 있는 칸이 다르면 이렇게 벌어진다. 반대로 전기·EMI는 두 수치가 "
            "같다 — 택일군 하나짜리 해석이라 채운 칸이 곧 준비된 재료다.")
    _p(doc, "(라) “적용 대비 준비”는 분모를 줄여 올린 수치다. 세라믹에 변형률-수명 피로계수를 "
            "요구하면 그 재료는 영원히 준비 안 됨으로 남는다. 알루미나로 strain-life 피로해석을 "
            "돌리는 일은 없으므로 분모에서 뺀다. 이 상승은 수집 성과가 아니라 정의다. "
            "그래서 두 열을 반드시 나란히 둔다.", bold=True)

    # ── 7. 검증 장치 ────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("7. 검증 장치 다섯 — 그리고 각각이 실제로 잡은 것", level=1)
    doc.add_heading("7.1  무결성 검사 — 구조를 본다", level=2)
    _tbl(doc, ["검사", "실제로 잡은 것"], [
        ["Mooney-Rivlin G0 음수",
         "Dragon Skin 10 MEDIUM(−86.0 kPa) · Ecoflex 00-50(−315.9 kPa). 초기 전단탄성률이 "
         "음수면 Drucker 안정조건 위반이라 해석이 발산한다. 4건 삭제 후 상시 항목으로 승격"],
        ["비금속인데 비열이 금속급",
         "EMC 236 J/(kg·K). 넣자마자 걸린 3건은 전부 오탐이었고 오탐의 원인이 진짜 결함이었다 — "
         "SmCo·Alnico·VACOFLUX의 category가 통째로 뒤섞여 있었다"],
        ["온도가 값인데 온도 조건",
         "융점·Tg·HDT 48건. 텅스텐 융점 3410 °C가 “상온에서 3387 °C 떨어진 값”으로 계산돼 "
         "대표값에서 맨 뒤로 밀렸다"],
        ["조건이 dict 아님",
         "SQLAlchemy JSON 컬럼이 파이썬 None을 SQL NULL이 아니라 JSON 문자열 'null'로 "
         "직렬화하고 있었다. none_as_null=True 한 줄로 두 경로가 동시에 고쳐졌다"],
        ["LCSR 가로축 비단조",
         "Kapton HN이 298 K에서 ID·TD 두 방향을 각각 3율속으로 재는데 온도만 보고 계열을 묶어 "
         "가로축이 중복되고 배율이 역행하는 *DEFINE_CURVE가 나가고 있었다"],
        ["Basquin 쌍 불일치",
         "VACOFLUX 50이 σf′를 σf′/E(무차원)로만 인쇄해 지수만 넣은 것을 결함으로 잡았다. "
         "conditions.pair_incomplete에 사유를 적은 행은 빼도록 고쳤다 — 사유 없는 반쪽만 잡는다"],
    ], bsz=7.5)
    doc.add_paragraph()

    doc.add_heading("7.2  반례 검증 — 선언은 주장이 아니라 반례로 무너진다", level=2)
    _p(doc, "“이 물성은 이 재료군에서 애초에 발표되지 않는다”는 선언은 두 조건을 다 만족할 때만 "
            "유효하다 — ① 문서에 기록된 사유가 있다(‘아직 못 찾았다’와 구분한다), "
            "② 그 군의 어느 재료도 그 물성을 실측으로 갖고 있지 않다(반례가 0이다). "
            "값이 하나라도 들어오면 그 쌍은 자동으로 부재에서 빠진다.", bold=True)
    _tbl(doc, ["사건", "결과"], [
        ["선언을 그대로 믿었을 때",
         "102칸이 “부재인데 값이 있음”으로 걸렸다. 벤조페논은 상용 화학물질이라 CRC에 밀도가 "
         "있고, laminate·coverlay 매처가 비열을 가진 재료까지 잡았다. 선언 16개 중 13개가 무너졌다"],
        ["반례에 tier 조건이 없었을 때",
         "8차가 아크릴 폼 테이프 11종에 항복강도 tier4 가정값을 넣자 “점착테이프에는 항복점이 "
         "물리적으로 없다”는 선언이 자동 은퇴했다. 반례는 실측(tier≤3)만 인정하도록 고쳤다 — "
         "우리가 만든 가정이 우리의 부재 선언을 지우면 안 된다"],
        ["페이월은 부재가 아니다",
         "접근 문제는 언제든 열릴 수 있으므로 분모에 남긴다(FR-4 방향별 열전도율, EMC 포아송비 등)"],
    ], bsz=7.5)
    doc.add_paragraph()

    doc.add_heading("7.3  자기 검산 — 얻은 값으로 그 논문의 다른 값을 예측한다", level=2)
    _tbl(doc, ["검산", "결과"], [
        ["Ti-6Al-4V Johnson-Cook 세트", "그 세트로 예측한 응력 2725 MPa 대 같은 논문 실측 1680 MPa. 폐기"],
        ["PC 비열 표", "그 논문의 PC 밀도가 1.44 g/cm³다(실제 1.20). 표 폐기"],
        ["FR-4 열물성 표",
         "같은 표의 Sn63Pb37 밀도가 1,890 kg/m³(실제 약 8,400). 알려진 재료로 표 전체를 검산해 "
         "통째로 버렸다"],
        ["EMC 비열 — 혼합물 상하한",
         "EMC는 에폭시(cp≈1,200)와 실리카(≈740)의 혼합물이라 1,200 초과나 740 미만은 불가능. "
         "논문 표 두 건(1,672 · 236)을 걸렀다. 둘 다 같은 표의 다른 행은 완벽했다 — 행 단위 판단"],
        ["Prony Σg",
         "Microsyst. Nanoeng. 본문 마스터커브가 g1+g2+g3 = 2.11로 상대계수가 성립하지 않았다. "
         "반대로 FR-4는 Σα = 0.9999·1.0022가 나와 그 열이 상대계수임을 확정했다"],
        ["Meggiolaro-Castro 추정식",
         "FD&E 실측 피팅 54건으로 자기 검산. σf′/Su 실측 중앙값 대 추정 — 탄소강 1.41 대 1.5, "
         "Ni계 1.56 대 1.4는 통과, 스테인리스 2.15 대 1.5는 불통. 편향 방향을 notes에 적고 "
         "SUS301은 실측 계열값으로 덮었다"],
        ["물리 한계로 단위 확정",
         "PORON 시트가 두 단위를 병기하고 값은 하나였는데, g/ft² 해석(36,600)이 37 °C 3 mm "
         "공기층의 순수확산 한계 약 32,800을 넘어 불가능하므로 g/m²로 확정"],
    ], bsz=7.5)
    doc.add_paragraph()

    doc.add_heading("7.4  원문 문자열 대조 — 저장한 숫자가 인용 출처에 실제로 있는가", level=2)
    _p(doc, "수집을 마친 뒤 각 행의 값이 인용 출처 원문에 문자열로 존재하는지 검사하는 스크립트를 "
            "돌린다. 5차 파동에서 한 배치가 자발적으로 만들었고 그 뒤 브리프의 작업 습관이 됐다.")
    _tbl(doc, ["잡은 것", "실체"], [
        ["Al2024-T3 표면에너지 35.3 mJ/m²", "값은 맞는데 엉뚱한 논문에 붙어 있었다"],
        ["Lexan FR25A 탄성계수", "SABIC 매뉴얼 출처에 붙었으나 실제는 arXiv 논문"],
        ["FR-4 밀도",
         "값은 g-10.pdf에 인쇄됐는데 블록 URL이 fr-4.pdf. 인제스트는 블록당 출처 하나만 "
         "저장하므로 그대로면 엉뚱한 문서에 붙는다"],
        ["아크릴 PSA 열전도율 0.176",
         "원문 전문에 0.176이 0회 등장. notes가 주장한 “0 wt% 대조군 실측”이 거짓이었다"],
        ["Copper Pillar Bump 항복강도 3건",
         "tier1 실측으로 있었으나 원문이 “theoretical yield strength”라고 밝힌 계산값"],
    ], bsz=7.5)
    doc.add_paragraph()
    _p(doc, "여기서 나온 교훈이 가장 중요하다. notes가 주장하는 근거를 원문에서 확인하지 않으면 "
            "틀린 값이 아니라 틀린 근거가 번진다. 값은 검산으로 걸리지만 근거는 안 걸린다.", bold=True)

    doc.add_heading("7.5  규칙 6 — 기존 값과의 자릿수 대조", level=2)
    _p(doc, "넣기 전에 DB에서 그 재료의 기존 값을 먼저 확인한다. 값 하나만 보면 완벽한데 그 "
            "재료의 카드 안에서 앞뒤가 안 맞는 경우를 잡는다.")

    doc.add_heading("7.6  검사가 오탐을 냈을 때의 규율", level=2)
    _p(doc, "세 번 있었고 대응이 매번 같았다 — 검사를 느슨하게 만들기 전에 오탐의 원인부터 본다.",
       bold=True)
    _tbl(doc, ["오탐", "원인", "고친 방식"], [
        ["가정값인데 tier4 아님(SmCo 3건)",
         "조건에서 부분문자열 assumption을 찾는데 direction: “bulk effective (isotropic "
         "assumption)” 같은 방법 서술에 걸렸다. 값은 초음파 실측",
         "구조화된 표지 \"assumption\": true만 보도록 좁혔다. 검사는 산문이 아니라 구조를 봐야 한다"],
        ["Prony 항번호 없음(36건)",
         "장기계수 E∞는 Prony 항이 아니라 평형값이다. 급수는 E(t)=E∞+ΣEᵢexp(−t/τᵢ)이고 "
         "E∞는 급수 밖이라 항번호가 없는 게 맞다", "검사에서 제외"],
        ["recaptcha 문자열로 실패 판정",
         "정상 페이지의 폼 스크립트에도 그 문자열이 있다. 한 배치가 caplinq 물성표 6건을 "
         "그렇게 잘못 버렸다", "판정을 “찾는 제품명이 본문에 실제로 있는가”로 바꿨다"],
    ], bsz=7.5)
    doc.add_paragraph()
    _p(doc, "반대 방향으로 가드가 좁아서 정상값을 버린 사례도 셋이다. 굴절률 가드 0.8~5.0이 "
            "c-Si(5.613 @400 nm)와 TiN(0.7946 @633 nm)을, 유전율 상한 1e5가 MnZn 페라이트의 "
            "정상값 14만을, Prony 완화시간 상한이 마스터커브 피팅값 1e14 s를 버렸다. "
            "가드는 자릿수 사고만 잡게 두고 물리 범위를 좁게 잡지 말 것.")

    # ── 8. 함정 지도 ────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("8. 함정 지도가 축적되는 방식 — 왜 실패를 남기는가", level=1)
    _tbl(doc, ["문서", "무엇을 담나", "목적"], [
        ["COLLECTION_BRIEF.md 6절", "판정 함정 — 어떤 표·사이트·표기를 조심할 것인가",
         "다음 배치가 같은 값을 다시 집지 않게"],
        ["BLOCKED_SOURCES.md", "접근 실패 — 유료/로그인/미공표/삭제/그림뿐",
         "같은 곳을 다시 뒤지지 않게"],
        ["PROPERTY_DATA_HISTORY.md", "경위 — 무엇을 시도했고 왜 틀렸나",
         "판단의 근거를 재구성할 수 있게"],
    ])
    doc.add_paragraph()
    _p(doc, "브리프 정본의 첫 문장이 이렇게 적혀 있다 — “이 문서의 값어치는 규칙이 아니라 "
            "6절의 함정 지도에 있다. 각 항목은 전부 실제 사고 기록이다.”", bold=True)
    doc.add_heading("8.1  실제로 시간을 아낀 증거", level=2)
    _tbl(doc, ["무엇이", "어떻게 아꼈나"], [
        ["검산 방법의 전파",
         "6차가 FR-4 표를 같은 표의 알려진 재료(Sn63Pb37 밀도 1,890)로 검산해 버렸고, 그 방법을 "
         "브리프 6절에 박았다. 8차가 같은 방식(혼합물 물리 상하한)으로 EMC 논문 표 두 건을 "
         "즉시 걸렀다 — 두 건 다 같은 표의 다른 행은 완벽했으므로 규칙이 없었으면 표 전체를 "
         "신뢰했을 것이다"],
        ["복사 계보 규칙",
         "PV 적층 물성표가 세 문헌에 바이트 단위로 복사돼 있음을 확인한 뒤로, 같은 숫자 조합이 "
         "여러 논문에 나오면 한 건만 검산하고 나머지를 같이 버린다. 독립 확인처럼 보이는 것을 "
         "세 번 검증하는 낭비가 사라졌다"],
        ["아카이브 회수분",
         "9차가 fde.uwaterloo.ca를 잃은 뒤 Wayback으로 26세트를 회수해 FDE_FATIGUE_ARCHIVE.md에 "
         "표로 남겼다. 문서 첫 줄이 “다음 파동은 이 표를 먼저 보고 필요한 것만 다시 받아라”이고, "
         "회수하지 못한 것까지 명시해 헛수고를 막았다"],
        ["광맥을 지도에 올린 것",
         "브리프 2절에 accudynetest polytable_03이 38종의 임계표면장력과 물 접촉각을 한 행에 "
         "나란히 인쇄한다고 적혀 있다. 9차 젖음성 배치가 182종 중 179종을 닫았고 그중 133종(74%)이 "
         "여기서 나왔다"],
        ["“성공처럼 보이는 실패” 목록",
         "hdl.handle.net이 862바이트 스텁을 준다는 항목이 없었다면 7차가 그 스텁을 근거로 "
         "“그 논문에 값이 없다”고 결론냈을 것이다"],
    ], bsz=7.5)
    doc.add_paragraph()
    _p(doc, "그리고 0건 보고를 실패로 읽지 않는다. 3차 파동의 0건 보고 세 건이 전부 값 있는 "
            "결과로 바뀌었다 — 점탄성 0건은 사유가 등급 불일치임을 밝혀 제약을 풀자 4종 115항이 "
            "나왔고, 무기 비열 0건은 CCL 벤더 6곳이 비열을 인쇄하지 않는다는 구조를 확정했으며, "
            "테이프 밀도 0건은 요청 밖에서 관찰한 12건으로 기존 DB 17건을 교차검증했다. "
            "어디를 그만 파야 하는지가 어디를 더 파야 하는지만큼 중요하다.", bold=True)

    # ── 9. 신뢰도 고도화 ────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("9. 신뢰도 고도화 경로 — 가정값을 실측으로 갈아 끼우는 구조", level=1)
    _tbl(doc, ["단", "무엇", "어떻게 위로 올라가나"], [
        ["tier4", "근거 있는 가정(계열 통계 중앙값 · 공표된 추정식)",
         "assumption:true + method:estimated로 새긴다. 실측이 들어오면 대표값 선택에서 자동으로 밀린다"],
        ["tier3", "계열 대표값 · 2차 인용 · 등급 불일치 대체값", "1차 출처를 찾으면 승격"],
        ["tier2", "핸드북 · 표준 · 공인 DB", "그 제품 실측을 찾으면 승격"],
        ["tier1", "그 제품 실측(벤더 TDS · 그 재료를 시험한 논문)", "종착"],
    ])
    doc.add_paragraph()
    _p(doc, "갈아 끼우는 데 사람 손이 필요 없다. 대표값 선택이 tier를 우선 순위로 쓰므로 tier1 "
            "행이 하나 들어오는 순간 조회·비교·DYNA 카드가 전부 그쪽을 쓴다. tier4 행은 삭제하지 "
            "않고 남긴다 — 불확실성 범위와 이력을 보존하기 위해서다.", bold=True)

    doc.add_heading("9.1  지금 tier4가 몇 건이고 어디에 몰려 있나", level=2)
    _t4 = tiers.get(4, 0)
    _p(doc, f"tier4는 {_t4:,}건({_t4*100/n_val:.1f}%)이다. 그중 assumption 표지를 가진 것이 "
            f"{n_assum:,}건이고, 나머지 {_t4-n_assum}건은 가정이 아니라 계산·유도값이다.")
    rows = []
    for k, tot_k, t4 in c.execute(
            "select property_key, count(*), sum(case when quality_tier=4 then 1 else 0 end) "
            "from property_value group by 1 having count(*)>=50 order by "
            "(1.0*sum(case when quality_tier=4 then 1 else 0 end)/count(*)) desc limit 12"):
        nm = c.execute("select name from property_definition where key=?", (k,)).fetchone()
        rows.append([(nm[0] if nm else k)[:26], f"{t4} / {tot_k}",
                     f"{t4*100/tot_k:.1f}%  {bar(t4*100/tot_k, 10)}"])
    _tbl(doc, ["물성(재료 50건 이상 보유)", "tier4 / 전체", "tier4 비율"], rows, bsz=7.5)
    doc.add_paragraph()
    _cat = c.execute("select m.category, count(*) from property_value pv join material m "
                     "on m.id=pv.material_id where pv.quality_tier=4 group by 1 "
                     "order by 2 desc").fetchall()
    _p(doc, "재료 카테고리별 tier4 — " + " · ".join(f"{k} {v:,}" for k, v in _cat) +
       ". 구멍이 거의 전부 폴리머·복합·고무·폼에 몰려 있고 금속·세라믹은 사실상 다 찼다.", bold=True)
    _tbl(doc, ["재료", "tier4 건수"],
         [[k[:56], v] for k, v in c.execute(
             "select m.name, count(*) from property_value pv join material m on m.id=pv.material_id "
             "where pv.quality_tier=4 group by 1 order by 2 desc limit 8")], bsz=7.5)
    doc.add_paragraph()

    doc.add_heading("9.2  포아송비 — 가정값의 실제 불확도가 측정된 사례", level=2)
    _p(doc, "포아송비 tier4가 여덟 개 숫자에서 복제돼 있었고 0.490 하나가 106건이었다. "
            "그런데 Isola는 모든 라미네이트 TDS에 ν를 length/cross로 ASTM D3039 실측으로 "
            "인쇄한다 — 실측 스펙트럼이 FR408HR 0.137 ~ I-Tera MT40 0.234다. "
            "“직물유리 적층판 ν=0.183”이라는 클래스 가정 하나가 20건에 복제돼 있는데 "
            "실제 폭이 0.13~0.23이다.", bold=True)
    _tbl(doc, ["뒤집힌 것", "실체"], [
        ["3M VHB 아크릴 폼 테이프의 ν",
         "0.49가 아니라 0.35다. 3M 공식 TDS 2018·2019 개정판에 Poisson's Ratio 0.35 행이 있는데 "
         "3M이 2022년 이후 개정판에서 이 행을 삭제했다. 폼 테이프 8종의 가정이 전부 틀렸다"],
        ["PA6-GF30의 등방 가정",
         "유동방향 0.43/0.39, 직교방향 0.22/0.27로 2배 방향 편차를 단일 가정값 0.35가 덮고 있었다"],
    ], bsz=7.5)
    doc.add_paragraph()
    _p(doc, "교훈은 tier4를 tier1로 갈아 끼우는 일이 “정밀도 개선”이 아니라 “틀린 값의 교체”인 "
            "경우가 있다는 것이다.", bold=True)

    # ── 10. 현재 상태 ───────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("10. 현재 상태와 남은 일", level=1)
    _cats = ", ".join(f"{k} {v}" for k, v in c.execute(
        "select category,count(*) from material group by 1 order by 2 desc"))
    _kinds = " · ".join(f"{k} {v}" for k, v in c.execute(
        "select kind,count(*) from source group by 1 order by 2 desc"))
    n_doi = c.execute("select count(*) from source where coalesce(doi,'')<>''").fetchone()[0]
    n_used = c.execute("select count(distinct property_key) from property_value").fetchone()[0]
    _tbl(doc, ["항목", "값"], [
        ["재료", f"{n_mat}종 ({_cats})"],
        ["물성값", f"{n_val:,}건"],
        ["출처", f"{n_src:,}건 (DOI 보유 {n_doi})"],
        ["물성 정의", f"{n_def}종 — 값이 있는 정의 {n_used}종(개통률 {n_used*100/n_def:.0f}%)"],
        ["출처 종류", _kinds],
    ])
    doc.add_paragraph()
    _tbl(doc, ["격자", "칸"], [
        ["전체 격자", f"{tot_cells:,}"],
        ["채워진 칸", f"{tot_filled:,}  ({tot_filled*100/tot_cells:.1f}%)"],
        ["그중 실측기반(tier≤3)", f"{tot_meas:,}  ({tot_meas*100/tot_cells:.1f}%)"],
        ["가정값이 채운 칸", f"{tot_filled-tot_meas:,}  "
                       f"({(tot_filled-tot_meas)*100/tot_cells:.1f}%p)"],
        ["빈칸", f"{tot_cells-tot_filled:,}"],
        ["그중 구조적 부재라 못 채우는 칸", f"{dead:,}"],
        ["실제로 채울 수 있는 남은 칸", f"{true_gap:,}"],
        ["(참고) 실측 기준으로 남은 일", f"{tot_cells-tot_meas:,}"],
    ])
    doc.add_paragraph()
    _tbl(doc, ["남은 덩어리", "칸", "왜 어려운가"],
         [[lever_label(k)[0], n, lever_label(k)[1]] for k, n in lever.most_common(8)], bsz=7.5)
    doc.add_paragraph()
    _p(doc, f"남은 일의 성격 요약 — {len(_only1)}개 해석은 재료마다 딱 한 칸씩 모자라고, "
            "낙하·충격만 2층 병목이다. 막힌 경로가 바뀌지 않는 한 남은 칸 전부는 안 닫힌다. "
            "두 파동에 걸쳐 확인된 결론은 율속과 피로가 OA 밖(핸드북·벤더 시험보고서·유료 "
            f"표준)에 있다는 것이다. 동시에, 채워진 {tot_filled:,}칸 중 {tot_filled-tot_meas:,}칸이 "
            "가정값이다 — 남은 칸을 닫는 것과 이미 닫은 칸의 근거를 두껍게 하는 것은 별개의 "
            "일이고 후자가 더 크다.", bold=True)

    # ── 11. 틀렸던 순간들 ───────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("11. 이 방법론이 틀렸던 순간들", level=1)
    _p(doc, "이 절이 이 문서에서 가장 중요하다. 아래는 전부 우리가 세운 가설·지표·지시가 반례로 "
            "무너진 기록이다. 자기 교정이 이 시스템의 핵심이고, 그 교정은 예외 없이 "
            "“숫자가 이상하다”에서 시작됐다.", bold=True)
    _tbl(doc, ["무엇을 믿었나", "실제", "고친 방식"], [
        ["544종 중 124종(22%)이 구조적 부재라 8%p가 천장 아래다",
         "재료 단위로 통째 제외해서 나온 숫자였다 — 이미 채워진 칸까지 버린 계산이다. "
         "부재를 (재료군 × 물성) 쌍으로 선언하고 반례를 붙이니 선언 16개 중 13개가 무너졌고 "
         "실제 부재는 106칸(0.7%)뿐이었다",
         "천장은 100%에 가깝다. 앞선 파동의 “확정된 부재”는 특정 제품을 못 찾았다는 뜻이지 "
         "그 재료군 전체를 뜻하지 않았다"],
        ["셀 채움 64.5%가 진척도다",
         "tier4를 빼면 46.3%였다. 18.1%p · 2,628칸이 가정값이었다. 구조적 부재(0.7%)를 찾으러 "
         "갔는데 그보다 26배 큰 왜곡을 찾았다",
         "지표에 실측기반 열을 넣어 둘을 나란히 보인다. 지표를 의심할 때는 분모만이 아니라 "
         "분자도 봐야 한다"],
        ["반례 검증 장치는 안전하다",
         "8차가 아크릴 폼 테이프에 tier4 가정값을 넣자 부재 선언이 자동 은퇴했다. 그 경고는 "
         "157종을 오탐하고 있었다",
         "반례를 실측(tier≤3)으로 한정. 방침을 바꾸면 그 방침이 기존 장치를 어떻게 흔드는지도 "
         "봐야 한다 — 결정 하나가 부재 판정·경고·커버리지 세 곳을 동시에 흔들었다"],
        ["Prony는 G부터 훑는 것이 최단 경로다",
         "145종 중 자기 제품의 G를 인쇄한 벤더는 0곳이고 G 단독으로 닫힌 것은 3종뿐이었다. "
         "반대로 Prony 세트는 포장·디스플레이 FEA 논문 표에 계속 실려 119종이 1,490행으로 닫혔다",
         "한 사례(VHB)에서 일반 규칙을 만든 것이 오류였다"],
        ["브리프에 적힌 접근 경로는 맞다",
         "2절이 accudynetest 색인을 polytable_01~04라 적어 왔는데 04는 404다. 여러 파동이 그 "
         "문장을 그대로 복사해 쓰고 있었다",
         "지시서도 1차 자료로 취급해 검증한다. 확인 과정에서 polytable_03이 두 키를 한 출처로 "
         "닫는 경로임을 찾았다"],
        ["근거 문구는 잘 저장되고 있다",
         "로더가 note(단수)만 읽는데 브리프 예시는 notes(복수)였다. 최근 인제스트분의 37%, "
         "전체 1,590건의 근거가 조용히 사라졌다. 값은 남고 “왜 그 값인가”만 없어져 무결성 "
         "검사에도 안 걸린다",
         "두 키를 다 받게 고치고 38개 파일을 재적재. 같은 종류가 세 번째다 — tier vs "
         "quality_tier(377건 강등), volume_resistivity vs resistivity_volume(없는 키를 셈)"],
        ["척도는 중립이다",
         "여섯 번 거짓말했다 — 없는 키를 셈 / 산소소광 0.0%(71종이 원리적으로 못 채우는 칸) / "
         "흡습에 24h만 요구 / 값이 잘못된 키에 들어가 안 보임 / 낙하 22%(146종 부당 미달) / "
         "tier를 구분 안 함",
         "물성을 가졌다는 것과 해석이 그걸 찾을 수 있다는 것은 다르다"],
        ["내가 주는 타당성 대역은 판정 기준이다",
         "EMC Csat 대역이 한 자릿수 낮았다. 에이전트가 그대로 믿었다면 유효한 데이터를 전부 "
         "잃었을 것이다. 검산 밴드도 재료군 중립이 아니다(Mg LCF는 c가 −1 근처가 정상)",
         "“내가 주는 대역은 참고이지 판정 기준이 아니다”를 브리핑에 넣었고, 이 한 줄이 "
         "에이전트의 자기 정정을 세 번 끌어냈다"],
        ["브리프의 우선순위는 목록과 맞다",
         "“1순위는 솔더”라고 썼는데 담당 목록 78종에 솔더가 한 종도 없었다. 유일한 솔더는 이미 "
         "피로 키를 다 갖고 있어 목록에 들어갈 수 없었다",
         "대상 목록에는 무엇이 없어서 뽑혔는지를 함께 적는다"],
        ["실패는 언제나 되돌릴 수 있다",
         "피로 배치가 fde.uwaterloo.ca를 무지연으로 610쪽 받은 뒤 IP 차단을 맞았다. "
         "AA7050-T7351 실측 피팅과 유일한 복합재 후보를 놓쳤고 둘 다 Wayback에도 없다",
         "이것만은 사후 교정이 불가능했다. “디렉터리 순회에 지연을 넣어라”를 6절에 박았지만 "
         "잃은 것은 돌아오지 않는다"],
    ], bsz=7)
    doc.add_paragraph()

    doc.add_heading("11.1  이 문서를 쓰면서 새로 발견한 것", level=2)
    _p(doc, f"(가) coverage_report.py의 “유효 채움”이 부재 칸을 이중으로 센다. 구조적 부재 "
            f"{htot['unfill']}칸 중 {htot['unfill_filled']}칸에는 이미 tier4 가정값이 들어 있다"
            "(8차 방침대로라 정상이다). 그런데 지표는 이 칸을 분자에는 채운 것으로 넣고 분모에서는 "
            "빼고 있어, 해석별 유효채움이 100%를 넘는다.", bold=True)
    _tbl(doc, ["해석", "부재 칸", "그중 이미 채워진 칸", "리포트의 유효채움", "부재 칸을 양쪽에서 뺀 값"],
         [[r["name"], r["unfill"], r["unfill_filled"],
           f"{r['filled']*100/r['eff']:.1f}%" if r["eff"] else "—",
           f"{r['eff_filled']*100/r['eff']:.1f}%" if r["eff"] else "—"]
          for r in hrows if r["unfill_filled"]], bsz=7.5)
    doc.add_paragraph()
    _p(doc, f"리포트가 인쇄하는 “남은 일은 {tot_cells-tot_filled}칸이 아니라 "
            f"{tot_cells-htot['unfill']-tot_filled}칸이다”도 같은 이유로 과소치다. "
            f"정직한 값은 {true_gap}칸이다({tot_cells-tot_filled} − {dead}, "
            f"{dead} = 부재이면서 아직 비어 있는 칸). 이 문서의 5·10절은 그 값을 쓴다.", bold=True)
    if odd:
        _p(doc, f"(나) 문서화된 method×tier 조합표가 지켜지지 않고 있고, 검사가 없다. "
                f"PROPERTY_DATA_HISTORY의 “정한 규칙” 2항이 허용 조합 7개를 표로 못박고 "
                f"“다른 조합이 나오면 어딘가 어긋난 것이다”라고 적는데, 현재 DB에 표 밖 조합이 "
                f"{sum(n for _,_,n in odd)}건({sum(n for _,_,n in odd)*100/n_val:.1f}%) 있다.",
           bold=True)
        _note = {("computed", 1): "논문이 상대계수와 기준값을 같은 표에 인쇄해 곱한 Prony 항 등 — 대체로 정당해 보인다",
                 ("computed", 2): "데이터시트 인쇄값의 단위 환산(열저항→열전도율, N/20mm→Pa 등)",
                 ("estimated", 1): "가장 문제다. tier1은 “그 제품의 인쇄된 실측”인데 method가 estimated다 — "
                                   "추정값이 tier1에 앉으면 실측이 들어와도 밀려나지 않는다",
                 ("estimated", 3): "계열 대입인데 tier3",
                 ("handbook", 1): "정성 등급 인용"}
        _tbl(doc, ["조합", "건수", "살펴본 결과"],
             [[f"{m} @ tier{t}", n, _note.get((m, t), "—")] for m, t, n in
              sorted(odd, key=lambda x: -x[2])], bsz=7.5)
        doc.add_paragraph()
        _p(doc, "특히 estimated @ tier1은 2절에서 설명한 자동 대체 장치를 정면으로 무력화한다. "
                "이 행들은 assumption 표지도 없어 기존 무결성 검사 2항목에 걸리지 않는다. "
                "integrity_check.py에 method×tier 조합 검사를 추가할 것을 권한다.", bold=True)

    doc.add_heading("11.2  요지", level=2)
    _p(doc, "이 절의 사례들에는 공통 구조가 있다. ① 숫자가 이상하다는 감각에서 시작한다. "
            "② 가설을 반증 가능한 형태로 만든다(부재를 재료군 × 물성 쌍으로 선언하고 반례를 센다). "
            "③ 반례가 하나라도 나오면 선언이 무너지게 자동화한다 — 사람이 기억할 필요가 없다. "
            "④ 그리고 무너진 기록을 남긴다 — 그래야 같은 가설을 다시 세우지 않는다.")
    _p(doc, "틀렸다는 것을 문서에 쓰는 것이 이 방법론의 산출물이다. 카탈로그의 신뢰도는 "
            "“우리가 옳았다”에서 오지 않고 “틀린 것을 어떻게 찾아냈는가”가 기록돼 있다는 "
            "데서 온다.", bold=True, size=11)

    # ── 부록 ────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("부록. 재현 명령과 파일 지도", level=1)
    _tbl(doc, ["목적", "명령"], [
        ["현재 커버리지", "python3 backend/scripts/catalog/coverage_report.py [--json]"],
        ["무결성 32항목 (0이 아니면 배포 금지)", "python3 backend/scripts/catalog/integrity_check.py"],
        ["파동 인제스트 (드라이런 → 적용)",
         "python3 backend/scripts/catalog/ingest_agent_json.py <디렉터리> [--apply]"],
        ["tier 정규화", "python3 backend/scripts/catalog/normalize_tiers.py [--apply]"],
        ["문서 재생성",
         "build_coverage_docx.py · build_methodology_docx.py · build_stats_xlsx.py"],
    ], bsz=7.5)
    doc.add_paragraph()
    _p(doc, "DB는 읽기만 한다. 쓰는 것은 ingest_agent_json.py --apply와 normalize_tiers.py "
            "--apply 두 경로뿐이고 둘 다 드라이런이 기본이다.", bold=True)
    _tbl(doc, ["파일", "역할"], [
        ["docs/COLLECTION_METHODOLOGY.md", "이 문서의 본문(마크다운)"],
        ["docs/COLLECTION_BRIEF.md",
         "수집 지시서 정본. 0절 원칙 · 0-A·0-B 방침 · 1절 tier · 2절 접근 경로 · "
         "3절 출력 형식 · 5절 작업 습관 · 6절 함정 지도"],
        ["docs/PROPERTY_DATA_HISTORY.md", "경위의 1차 사료. 72개 장"],
        ["docs/BLOCKED_SOURCES.md", "막힌 경로 대장 + 뚫렸던 경로"],
        ["docs/FDE_FATIGUE_ARCHIVE.md", "사이트를 잃고 Wayback에서 회수한 피로계수 26세트"],
        ["backend/app/property_taxonomy.py", f"물성 정의 {n_def}종"],
        ["backend/scripts/catalog/coverage_report.py", "지표 정의 · 구조적 부재 선언 · 반례 검증"],
        ["backend/scripts/catalog/integrity_check.py", "무결성 32항목 + 주의 1항목"],
        ["backend/scripts/catalog/ingest_agent_json.py", "적재 6단계 검증 · 조건 충돌 감지"],
        ["backend/scripts/catalog/normalize_tiers.py", "tier를 “값의 성격”으로 재판정"],
    ], bsz=7.5)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"저장: {OUT}")
    print(f"  해석 {len(cov)}개 · 격자 {tot_cells:,}칸 · 셀 채움 {tot_filled*100/tot_cells:.1f}% · "
          f"실측기반 {tot_meas*100/tot_cells:.1f}%")
    print(f"  채울 수 있는 남은 칸 {true_gap}칸 · 1칸 부족 재료 "
          f"{sum(1 for v in short_all.values() if v == 1)}종 · "
          f"표 밖 method×tier {sum(n for _,_,n in odd)}건")


if __name__ == "__main__":
    main()
